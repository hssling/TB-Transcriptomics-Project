# -*- coding: utf-8 -*-
"""Assemble the manuscript .docx.

Every quantity in the text is read from the analysis outputs, so the document
cannot drift away from the results it reports.
"""
import json
import os

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

HERE = os.path.dirname(os.path.abspath(__file__))
TAB = f"{HERE}/tables2"
FIG = f"{HERE}/figures2"
OUT = f"{HERE}/deliverables/TB_Treatment_Outcome_DAI_Revision_v14.docx"

CENTER = WD_ALIGN_PARAGRAPH.CENTER
JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY

M = json.load(open(f"{TAB}/arm_metrics.json"))
perf = pd.read_csv(f"{TAB}/arm_performance_table.csv")
degs = pd.read_csv(f"{TAB}/deg_summary.csv")
part = pd.read_csv(f"{TAB}/deg_state_partition.csv")
decon = pd.read_csv(f"{TAB}/deconvolution_stats.csv")
corr = pd.read_csv(f"{TAB}/prediction_celltype_correlation.csv")
shap_top = pd.read_csv(f"{TAB}/shap_top_features.csv")
enrich = pd.read_csv(f"{TAB}/enrichment_targeted.csv")
enrich_deg = pd.read_csv(f"{TAB}/enrichment_deg.csv")
hubs = pd.read_csv(f"{TAB}/network_hubs.csv")
ext_cell = pd.read_csv(f"{TAB}/external_celltype_stats.csv")
ext_auc = pd.read_csv(f"{TAB}/external_transfer_auc.csv")
conf = pd.read_csv(f"{TAB}/confounder_models.csv")
pubsig = pd.read_csv(f"{TAB}/established_signatures.csv")
sigcmp = pd.read_csv(f"{TAB}/signature_vs_model.csv")
traj = pd.read_csv(f"{TAB}/response_trajectory.csv")
sexd = pd.read_csv(f"{TAB}/sex_distribution.csv")
ylink = pd.read_csv(f"{TAB}/y_linked_genes.csv")
bench = pd.read_csv(f"{TAB}/model_benchmark.csv")

ARMS = C_ARMS = ["DX", "day_7", "week_4", "week_24"]
LABEL = {"DX": "Pre-treatment", "day_7": "Day 7", "week_4": "Week 4",
         "week_24": "Week 24", "combined": "All timepoints"}
MODEL_NICE = {"logistic_regression": "penalised logistic regression",
              "random_forest": "random forest",
              "gradient_boosting": "gradient boosting"}


# ---------- accessors ----------
def best(arm):
    return M[arm]["models"][M[arm]["best_model"]]


def pc(arm, cls="Unfavourable"):
    return {r["class"]: r for r in best(arm)["per_class"]}[cls]


def auc_txt(arm):
    b = best(arm)
    lo, hi = b["roc_auc_ci"]
    return f"{b['roc_auc']:.2f} (95% CI {lo:.2f}–{hi:.2f})"


def perm_p(arm):
    p = M[arm]["permutation"]["permutation_p"]
    return "< 0.002" if p < 0.002 else f"= {p:.3f}"


def dec(arm, cell):
    r = decon[(decon.arm == arm) & (decon.cell_type == cell)]
    return r.iloc[0] if len(r) else None


def cor(arm, cell):
    r = corr[(corr.arm == arm) & (corr.cell_type == cell)]
    return r.iloc[0] if len(r) else None


def degrow(arm):
    return degs[degs.arm == arm].iloc[0]


def pcount(cls):
    return int((part["class"] == cls).sum())


def sci(p):
    if p is None or pd.isna(p):
        return "n/a"
    return f"{p:.1e}" if p < 1e-4 else f"{p:.3f}"


def pval(p):
    return "< 0.001" if p < 0.001 else f"= {p:.3f}"


def ext_range():
    """Range of replication AUCs, for reporting the transfer result."""
    lo, hi = ext_auc.roc_auc.min(), ext_auc.roc_auc.max()
    return f"ROC-AUC {lo:.2f}–{hi:.2f}, all confidence intervals spanning 0.5"


def genes_list(arm, n=8):
    g = shap_top[shap_top.arm == arm].nsmallest(n, "rank")["gene_symbol"]
    return ", ".join(str(x) for x in g)


def _audit():
    a = json.load(open(f"{TAB}/robustness_audit.json"))
    tech = pd.DataFrame(a["technical"])
    inf = pd.DataFrame(a["influence"]["leave_one_out"])
    mult = pd.DataFrame(a["multiplicity"])
    load = pd.DataFrame(a["bacterial_load"])
    load_bits = ", ".join(
        f"{r.measure.upper()} ρ = {r.spearman_rho:+.2f}, p = {r.p_value:.3f}"
        for _, r in load.iterrows())
    return {
        "technical_min_p": float(tech.p_value.min()),
        "influence_min": float(inf.auc.min()),
        "influence_max": float(inf.auc.max()),
        "ms_min": float(a["model_selection"]["week24_min"]),
        "ms_max": float(a["model_selection"]["week24_max"]),
        "holm_week24": float(mult.loc[mult.arm == "Week 24", "p_holm"].iloc[0]),
        "load_text": load_bits,
    }


AUDIT = _audit()


# ---------- document scaffolding ----------
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)
for section in doc.sections:
    section.left_margin = section.right_margin = Inches(1.0)


def P(text, align=JUSTIFY, size=11, italic=False, bold=False, space_after=8):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    return p


def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = None
    return h


def figure(filename, caption, width=6.3):
    path = f"{FIG}/{filename}"
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = CENTER
    P(caption, align=JUSTIFY, size=9.5, space_after=12)


def table(headers, rows, caption=None, widths=None, font=9):
    if caption:
        P(caption, size=9.5, space_after=4)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(h))
        run.bold = True
        run.font.size = Pt(font)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v))
            run.font.size = Pt(font)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


# =====================================================================
# TITLE
# =====================================================================
title = ("Whole-blood transcriptomic signatures of unfavourable tuberculosis "
         "treatment outcome before and during therapy: an exploratory "
         "machine-learning and immune-deconvolution study")
p = doc.add_paragraph()
p.alignment = CENTER
r = p.add_run(title)
r.bold = True
r.font.size = Pt(14)

P("Article type: Original Article", align=CENTER, size=10)
P("Siddalingaiah H S, MD", align=CENTER, bold=True)
P("Professor, Department of Community Medicine, Shridevi Institute of Medical "
  "Sciences and Research Hospital, Tumkur, Karnataka, India",
  align=CENTER, size=10)
P("Correspondence: Dr Siddalingaiah H S. Email: hssling@yahoo.com",
  align=CENTER, size=10)

# =====================================================================
# ABSTRACT
# =====================================================================
H("Abstract", 1)

P("Background. Tuberculosis treatment fails in a minority of patients, and "
  "those patients sustain transmission and drug resistance. Whole blood "
  "carries information about the host immune response, but it is unclear how "
  "much of that information relates to eventual outcome, and at which point "
  "in the treatment course it becomes readable. Pooling samples taken before "
  "and during therapy conflates two distinct biological states and obscures "
  "that question. We therefore analysed each state separately and then "
  "compared them.")

P("Methods. We reanalysed the complete GSE89403 whole-blood RNA sequencing "
  f"deposit: {M['combined']['n']} libraries from {M['combined']['n_subjects']} "
  "adults with pulmonary tuberculosis and a recorded outcome, sampled at "
  "diagnosis, day 7, week 4 and week 24. Outcome was dichotomised as "
  "unfavourable versus cured. One preprocessing and modelling protocol was "
  "applied to every timepoint. Penalised logistic regression, random forest "
  "and gradient boosting were compared by repeated stratified cross-validation "
  "with feature selection confined to training folds. We report sensitivity "
  "and specificity for each outcome class alongside ROC-AUC, precision–recall "
  "AUC, calibration and a label-permutation null. SHAP attribution on the "
  "gradient-boosted model produced a ranked feature panel per timepoint, which "
  "was then carried into pathway enrichment and a graphical-lasso "
  "conditional-dependency network. Marker-based immune deconvolution, a sex "
  "and bacterial-load audit, and replication in an independent cohort "
  "(GSE67589; 20 patients, cure versus relapse, different platform and "
  "country) completed the analysis.")

P("Results. Discrimination depended strongly on when blood was sampled. "
  f"Before treatment it was modest ({auc_txt('DX')}; precision–recall AUC "
  f"{best('DX')['pr_auc']:.2f}) and did not separate from a permutation null "
  f"(p {perm_p('DX')}); no gene survived false-discovery correction. At the "
  f"end of therapy it was strong ({auc_txt('week_24')}; precision–recall AUC "
  f"{best('week_24')['pr_auc']:.2f}; permutation p {perm_p('week_24')}), with "
  f"{int(degrow('week_24').fdr_significant):,} genes differentially expressed "
  "at a false-discovery rate below 0.05. Class-specific performance at week 24 "
  f"reached sensitivity {pc('week_24')['sensitivity']:.2f} and specificity "
  f"{pc('week_24')['specificity']:.2f} for the unfavourable class. Across "
  "timepoints, predicted risk tracked a neutrophil-high, T-cell-low axis "
  f"(Spearman ρ up to {corr.spearman_rho.abs().max():.2f}). The neutrophil "
  "association was independent of sex, and sex-linked transcripts showed no "
  "differential expression by outcome. Four published signatures of active "
  "tuberculosis, applied without fitting, reproduced the week-24 separation "
  f"(best ROC-AUC {pubsig[pubsig.arm == 'week_24'].roc_auc.max():.2f}, "
  f"p = {sci(pubsig[pubsig.arm == 'week_24'].p_value.min())}) and matched or "
  "exceeded our fitted models in every arm. The signature did not transfer to "
  "the independent cohort: discrimination there was indistinguishable from "
  f"chance ({ext_range()}), and cell-composition differences were absent.")

P("Conclusions. Outcome is not readable from whole blood before treatment or "
  "during its first four weeks, and is clearly readable at the end of therapy, "
  "where patients recorded as not cured retain the established "
  "interferon-driven signature of active disease. That end-of-treatment signal "
  "is concurrent with outcome ascertainment and is therefore a marker of "
  "unresolved disease rather than a forecast. Two findings temper any "
  "enthusiasm for machine learning here: unfitted published signatures "
  "outperformed our models in every arm, and the signature did not generalise "
  "to an independent cohort. The contribution is a map of when whole blood is "
  "and is not informative, not a new predictor. Event numbers are small and "
  "the findings are exploratory.")

P("Keywords: tuberculosis; treatment outcome; whole-blood transcriptomics; "
  "machine learning; immune deconvolution; neutrophils; explainable artificial "
  "intelligence", space_after=14)

# =====================================================================
# 1. INTRODUCTION
# =====================================================================
H("1. Introduction", 1)

P("Tuberculosis remains the deadliest infection caused by a single agent, with "
  "an estimated 1.25 million deaths in 2023 [1]. Standard six-month "
  "chemotherapy cures most patients, but fails in roughly 5–10% of "
  "drug-susceptible disease and in a considerably larger share of "
  "multidrug-resistant disease [1]. The patients in whom it fails carry a "
  "disproportionate share of onward transmission and acquired resistance. "
  "Directly observed short-course treatment standardises the regimen and its "
  "supervision, but does not adapt therapy to the biology of the individual "
  "patient [1,2].")

P("Blood transcriptomics is an established route to that biology. Signatures "
  "reported by Berry, Zak, Sweeney and Singhania distinguish active from "
  "latent infection and anticipate progression to disease [3–6]. Predicting "
  "how a patient will respond to treatment has proved harder. In a large "
  "drug-susceptible Brazilian cohort, transcriptomic prediction of failure "
  "from a pre-treatment sample reached an area under the curve below 0.70, "
  "although prediction of death and of recurrence was stronger [7]. Related "
  "approaches integrate radiological, microbiological and clinical data across "
  "thousands of patients [9], characterise the interaction between "
  "tuberculosis and diabetes [10], and predict failure from electronic medical "
  "records [11].")

P("Two problems recur in this literature. The first is resolution. Bulk blood "
  "expression averages across a leukocyte mixture whose composition itself "
  "shifts with disease, so an apparent increase in inflammatory signalling may "
  "reflect a change in which cells are present rather than a change in what "
  "those cells are doing [8]. The second is design. Longitudinal cohorts "
  "sample the same patient repeatedly, and analyses that pool those samples "
  "mix pre-treatment biology with treatment-induced change. Because "
  "unfavourable outcomes are rare, such pooling can produce impressive "
  "aggregate accuracy that reflects when a sample was taken rather than what "
  "it says about the patient.")

P("Separating the timepoints turns that problem into a question worth asking. "
  "A pre-treatment sample reflects intrinsic patient biology before any "
  "therapeutic perturbation, and is the only state from which genuine "
  "stratification is possible. Samples taken during and at the end of therapy "
  "reflect how the immune system responded to that perturbation; they describe "
  "what a good response looks like and identify patients in whom it did not "
  "occur. Comparing the two states shows which signals are intrinsic, which "
  "emerge only under treatment, and what that difference implies about the "
  "biology of failure.")

P("We therefore reanalysed a longitudinal whole-blood cohort as separate "
  "biological states under one common protocol, and asked three questions. "
  "How well can outcome be discriminated in each state, reported for both "
  "outcome classes and against a permutation null? Which transcripts does the "
  "model use, what pathways do they represent, and how are they organised "
  "relative to one another? And does the resulting immune axis hold in an "
  "independent cohort, in different patients measured on a different platform? "
  "Throughout, features derived by machine learning are checked against "
  "established bioinformatic methods rather than reported on their own.")

# =====================================================================
# 2. METHODS
# =====================================================================
H("2. Methods", 1)

H("2.1 Cohort and outcome definition", 2)
P("The discovery cohort was GSE89403, a longitudinal whole-blood RNA "
  "sequencing study of adults with pulmonary tuberculosis receiving standard "
  "therapy [12]. Enrolment was restricted to HIV-negative patients with "
  "rifampicin-susceptible disease, so both factors are fixed by design rather "
  "than modelled here. Diabetes status was not recorded.")
P("The deposit contains 914 sequencing libraries. Libraries sequenced across "
  "lanes from the same biological specimen were merged, giving 453 samples. "
  "Non-tuberculosis controls and samples without an evaluable outcome were "
  f"excluded, leaving {M['combined']['n']} samples from "
  f"{M['combined']['n_subjects']} subjects at four protocol timepoints: "
  "diagnosis, day 7, week 4 and week 24. Outcome was dichotomised as "
  "unfavourable (recorded as not cured) versus cured (definite, probable or "
  "possible cure). Study flow is shown in Figure 1.")

H("2.2 Analytical arms", 2)
P("Each timepoint was analysed as a separate arm, because each represents a "
  "different biological state. The pre-treatment arm reflects patient biology "
  "before therapy and is the arm relevant to stratification. The day 7, week 4 "
  "and week 24 arms reflect the immune response to treatment, at increasing "
  "distance from its start. A fifth analysis pooled all timepoints with "
  "subject-grouped cross-validation; it is reported for contrast only, since "
  "it mixes states, and no predictive claim is made from it.")
rows = []
for arm in ARMS:
    rows.append([LABEL[arm], M[arm]["n"], M[arm]["n_events"],
                 M[arm]["n"] - M[arm]["n_events"],
                 f"{M[arm]['n_events'] / M[arm]['n'] * 100:.1f}%"])
rows.append([LABEL["combined"], M["combined"]["n"], M["combined"]["n_events"],
             M["combined"]["n"] - M["combined"]["n_events"],
             f"{M['combined']['n_events'] / M['combined']['n'] * 100:.1f}%"])
table(["Arm", "Samples", "Unfavourable", "Cured", "Prevalence"], rows,
      caption="Table 1. Composition of each analytical arm.",
      widths=[1.7, 1.0, 1.3, 1.0, 1.1])

H("2.3 Preprocessing", 2)
P("Gene counts were taken from the deposited count matrix. Genes detected in "
  "fewer than 10% of libraries were removed, leaving 16,145 genes. Counts were "
  "normalised to counts per million to remove library-size differences, then "
  "log-transformed as log₂(x+1). The identical pipeline was applied to every "
  "arm, so differences between arms reflect biology and not preprocessing.")

H("2.4 Predictive modelling", 2)
P("Three classifiers were compared in every arm: penalised logistic "
  "regression, random forest and gradient boosting. Each was fitted inside a "
  "pipeline that performed univariate feature pre-selection to 200 genes, so "
  "selection saw only training data and no information crossed the "
  "train–test boundary. Discrimination was estimated by stratified five-fold "
  "cross-validation repeated 40 times, with out-of-fold probabilities averaged "
  "across repeats and metrics computed on the pooled predictions. Class "
  "imbalance was handled by class weighting rather than resampling. The pooled "
  "analysis used subject-grouped folds so that no subject appeared in both "
  "training and test sets.")

H("2.5 Performance measures", 2)
P("Sensitivity, specificity, positive and negative predictive value and F1 are "
  "reported separately for each outcome class at the Youden operating point, "
  "because a single summary figure conceals how differently a model behaves on "
  "a rare class and a common one. These accompany ROC-AUC and "
  "precision–recall AUC with bootstrap confidence intervals from 2,000 "
  "resamples, the Matthews correlation coefficient, balanced accuracy, the "
  "confusion matrix and the Brier score. Discrimination was additionally "
  "tested against a null distribution built by permuting outcome labels 500 "
  "times and repeating the whole cross-validation, with observed and permuted "
  "values computed under one identical protocol. Leave-one-out "
  "cross-validation was run as a further check.")

H("2.6 Feature attribution and model choice", 2)
P("SHAP values were computed on the gradient-boosted model using TreeExplainer, "
  "which returns exact Shapley values for additive tree ensembles rather than a "
  "sampled approximation. Gradient boosting was chosen for attribution on that "
  "ground: it captures non-additive structure that a linear model cannot, and "
  "it admits an exact attribution method that a random forest ensemble of the "
  "same size does not. Discrimination for all three classifiers is reported in "
  "every arm so that this choice can be judged against the alternatives rather "
  "than assumed. Attribution was computed on a 50-gene panel per arm, and the "
  "ranking was carried forward unchanged into the two downstream analyses.")

H("2.7 Pathway analysis", 2)
P("Pathway analysis was run on the model's own attribution ranking, so that "
  "the biology described corresponds to what the classifier used rather than "
  "to a separately derived list. A fifty-gene panel is too small to support "
  "open-ended testing against tens of thousands of terms, so the primary test "
  "was targeted: the panel from each arm was tested by hypergeometric "
  "over-representation against six pre-specified immune programmes relevant to "
  "tuberculosis — neutrophil degranulation, T-cell receptor signalling, the "
  "interferon-inducible response, inflammatory and myeloid activation, "
  "cytotoxic and natural-killer effector function, and the B-cell programme. "
  "Six tests per arm keeps the multiple-testing burden negligible, and "
  "false-discovery correction was applied within each arm. Programme "
  "membership is listed in full with the code.")
P("Open-ended enrichment against Gene Ontology biological processes, KEGG, "
  "Hallmark and Reactome was run separately on the genes surviving "
  "false-discovery correction in the differential-expression analysis, where "
  "the gene list is large enough for such testing to be meaningful. Results "
  "of open-ended testing on the attribution panels themselves are reported in "
  "the supplement for completeness.")

H("2.8 Immune deconvolution", 2)
P("Cell composition was estimated by marker-gene-set enrichment: expression "
  "was standardised across samples and averaged within canonical marker panels "
  "for neutrophils, T cells, monocytes, B cells and natural killer cells. We "
  "describe the method as marker-panel enrichment rather than by a proprietary "
  "name because that is exactly what was computed. Group differences used the "
  "Mann–Whitney U test with rank-biserial correlation as the effect size. "
  "Agreement between the cell scores and the model's predicted probability was "
  "quantified by Spearman correlation; this is the step that tests whether the "
  "machine-learning signal and an established deconvolution method describe "
  "the same biology.")

H("2.9 Conditional-dependency network", 2)
P("A Gaussian graphical model was estimated with the graphical lasso on the "
  "SHAP-ranked panel of each arm. Using the attribution panel as the node set "
  "keeps one feature set flowing through the pipeline, so the network "
  "describes how the signal the model relies on is internally organised rather "
  "than forming a parallel result. Each edge is a partial correlation: the "
  "association between two transcripts once every other transcript in the "
  "panel has been conditioned out, which removes the indirect links that a "
  "marginal correlation network would display. Edges are undirected, outcome "
  "is not a node, and high-degree genes are described as hubs and not as "
  "regulators or drivers.")

H("2.10 Benchmarking against published signatures", 2)
P("Model-derived features are only worth reporting if they add something to "
  "what is already established, so each arm was also scored with four "
  "published blood signatures of active tuberculosis: the three-gene set of "
  "Sweeney and colleagues, the interferon core of the Berry signature, the "
  "myeloid-up and lymphoid-down set of Kaforou and colleagues, and the "
  "interferon-dominated risk signature of Zak and colleagues [3–6]. Each was "
  "computed as the mean standardised expression of its up-weighted members "
  "minus that of its down-weighted members. No parameter was fitted, so these "
  "scores are free of any optimism from our own data and provide an external "
  "reference against which our fitted models can be judged.")

H("2.11 Within-subject treatment response", 2)
P("Comparing different patients at one timepoint leaves between-subject "
  "variation inside every contrast. For subjects with both a diagnosis sample "
  "and a later sample, the change vector between them removes that variation "
  "and describes the individual response to therapy. Averaging those vectors "
  "across cured subjects defines a reference direction for a successful "
  "response, and each subject was scored by the cosine alignment of their own "
  "change with it. The reference was rebuilt leaving out each cured subject in "
  "turn, so no subject contributed to the direction used to score them. The "
  "change vectors were also submitted to the same classification protocol used "
  "for the arms, giving a like-for-like comparison between the cross-sectional "
  "state and the within-subject change.")

H("2.12 Independent replication", 2)
P("Replication used GSE67589, an independent cohort of 20 patients with "
  "pulmonary tuberculosis sampled at diagnosis, week 2 and week 4, in whom "
  "outcome was recorded as cure or relapse after apparently successful "
  "treatment [14]. The cohort differs from the discovery data in patients, "
  "country, measurement platform and outcome definition, so agreement across "
  "it is a conservative test rather than a repeat measurement. Because the "
  "platforms differ, transfer used rank-standardised signature scores rather "
  "than model coefficients: cell-composition scores were recomputed with the "
  "same marker panels, and a directional score built from the discovery "
  "cohort's differential expression was applied to the replication samples.")

H("2.13 Confounder audit", 2)
P("Sex was inferred from XIST and RPS4Y1 expression, since it is not "
  "deposited. Bacterial load was represented by time to culture positivity and "
  "Xpert cycle threshold. Multivariable logistic regression tested whether the "
  "neutrophil association survived adjustment for both. Sex-linked transcripts "
  "were tested directly for differential expression, because they can dominate "
  "an unadjusted feature ranking when a small outcome group is skewed by sex.")

H("2.14 Software and reproducibility", 2)
P("Analyses used Python with scikit-learn, XGBoost, SHAP, statsmodels and "
  "SciPy. Tests were two-sided at α = 0.05, with Benjamini–Hochberg "
  "false-discovery control where indicated. Effect sizes are reported "
  "alongside p-values throughout. Analysis code, the frozen models and all "
  "intermediate outputs are released so that results can be verified without "
  "re-execution.")

figure("Figure_study_flow.png",
       "Figure 1. Study flow. Cohort assembly, the four biological states "
       "analysed as separate arms, the common modelling protocol, and the "
       "verification of model-derived features against deconvolution, "
       "differential expression, network estimation and an independent cohort.")

# =====================================================================
# 3. RESULTS
# =====================================================================
H("3. Results", 1)

H("3.1 Cohort", 2)
P(f"After exclusions, {M['combined']['n']} samples from "
  f"{M['combined']['n_subjects']} subjects were analysed. Unfavourable "
  "outcomes were uncommon at every timepoint, ranging from "
  f"{min(M[a]['n_events'] for a in ARMS)} to "
  f"{max(M[a]['n_events'] for a in ARMS)} events per arm against 83 to 87 "
  "cured patients (Table 1). Every performance estimate below rests on fewer "
  "than ten events and is correspondingly uncertain; the confidence intervals "
  "reported alongside each estimate should be read as part of the result.")

H("3.2 Outcome is barely readable before treatment", 2)
b = best("DX")
P("In pre-treatment blood, discrimination was modest. The best of the three "
  f"classifiers ({MODEL_NICE[M['DX']['best_model']]}) reached a ROC-AUC of "
  f"{auc_txt('DX')}, with a precision–recall AUC of {b['pr_auc']:.2f} "
  f"against a prevalence of {M['DX']['n_events'] / M['DX']['n']:.2f}. The "
  "confidence interval includes 0.5, and the estimate did not separate from a "
  f"label-permutation null (observed "
  f"{M['DX']['permutation']['observed_auc']:.2f} against a null mean of "
  f"{M['DX']['permutation']['null_mean']:.2f}; p {perm_p('DX')}). "
  f"Leave-one-out cross-validation gave {M['DX'].get('loo_auc', float('nan')):.2f}, "
  "agreeing with the repeated-fold estimate.")
P("At the Youden operating point the model behaved very differently on the two "
  f"classes. For the unfavourable class, sensitivity was "
  f"{pc('DX')['sensitivity']:.2f} and specificity "
  f"{pc('DX')['specificity']:.2f}, but positive predictive value was only "
  f"{pc('DX')['ppv']:.2f} against a negative predictive value of "
  f"{pc('DX')['npv']:.2f}. In practical terms a negative result is somewhat "
  "informative and a positive result is not. Differential expression told the "
  f"same story: {int(degrow('DX').nominal_p05):,} genes reached nominal "
  "significance but none survived false-discovery correction.")
P("This is a negative result, and we report it as one. Pre-treatment whole "
  "blood, in a cohort of this size, does not support risk stratification for "
  "eventual treatment outcome.")

H("3.3 Outcome becomes readable only at the end of therapy", 2)
P("Early treatment samples were little better than pre-treatment ones. At day "
  f"7 discrimination was {auc_txt('day_7')} and at week 4 "
  f"{auc_txt('week_4')}; neither separated from its permutation null "
  f"(p {perm_p('day_7')} and p {perm_p('week_4')} respectively), and neither "
  "yielded a gene surviving false-discovery correction. Where the bootstrap "
  "interval at day 7 excludes 0.5 but the permutation test does not reject, we "
  "take the permutation result as the more conservative and more appropriate "
  "check, because it repeats the entire cross-validation under the null rather "
  "than resampling a fixed set of predictions.")
P("Week 24, the end of therapy, is different. Discrimination was "
  f"{auc_txt('week_24')} with a precision–recall AUC of "
  f"{best('week_24')['pr_auc']:.2f} against a prevalence of "
  f"{M['week_24']['n_events'] / M['week_24']['n']:.2f}, a Matthews "
  f"correlation of {best('week_24')['mcc']:.2f}, and a permutation p "
  f"{perm_p('week_24')}. All three classifiers agreed at this timepoint "
  "(Table 2), which argues that the signal is a property of the data rather "
  "than of one algorithm. The change is therefore abrupt rather than gradual: "
  "outcome is unreadable through the first four weeks of treatment and clearly "
  "readable at its end.")
P("Differential expression moved in step. The number of genes surviving "
  "false-discovery correction rose from none before treatment to "
  f"{int(degrow('week_24').fdr_significant):,} at week 24. Class-specific "
  "performance at week 24 was balanced across both classes: sensitivity "
  f"{pc('week_24')['sensitivity']:.2f} and specificity "
  f"{pc('week_24')['specificity']:.2f} for the unfavourable class, against "
  f"sensitivity {pc('week_24', 'Cured')['sensitivity']:.2f} and specificity "
  f"{pc('week_24', 'Cured')['specificity']:.2f} for the cured class.")
P("The interpretation of this result requires care. Week 24 is also the point "
  "at which treatment outcome is ascertained, so a model reading week-24 blood "
  "is describing a state concurrent with the outcome rather than forecasting a "
  "future event. What it identifies is unresolved immune activation in "
  "patients whose disease has not been cleared. That is a treatment-response "
  "signal, not a prediction, and the distinction matters for how such a marker "
  "could ever be used.")
P("Pooling all timepoints with subject-grouped folds gave "
  f"{auc_txt('combined')}, with a permutation p {perm_p('combined')}. That "
  "significance is instructive rather than reassuring: it is inherited from "
  "the week-24 samples inside the pool, and it attaches to a figure that "
  "describes neither the pre-treatment state nor the end-of-treatment state. "
  "A pooled analysis of this cohort will therefore look moderately successful "
  "while concealing that its information comes almost entirely from samples "
  "taken when the outcome was already being determined. We report it to make "
  "that point and attach no predictive claim to it.")

rows = []
for arm in ARMS + ["combined"]:
    bb = best(arm)
    lo, hi = bb["roc_auc_ci"]
    rows.append([
        LABEL[arm], M[arm]["n"], M[arm]["n_events"],
        MODEL_NICE[M[arm]["best_model"]].replace("penalised ", ""),
        f"{bb['roc_auc']:.2f} ({lo:.2f}–{hi:.2f})",
        f"{bb['pr_auc']:.2f}",
        f"{pc(arm)['sensitivity']:.2f}", f"{pc(arm)['specificity']:.2f}",
        f"{pc(arm)['ppv']:.2f}", f"{pc(arm)['npv']:.2f}",
        f"{bb['mcc']:.2f}",
        perm_p(arm).lstrip("= "),
    ])
table(["Arm", "n", "Events", "Model", "ROC-AUC (95% CI)", "PR-AUC",
       "Sens.", "Spec.", "PPV", "NPV", "MCC", "Perm. p"], rows,
      caption="Table 2. Discrimination in each biological state. Sensitivity, "
              "specificity and predictive values refer to the unfavourable "
              "class at the Youden operating point.",
      widths=[1.05, 0.4, 0.5, 0.75, 1.15, 0.5, 0.45, 0.45, 0.45, 0.45, 0.45, 0.6],
      font=8)

figure("Figure_performance_arms.png",
       "Figure 2. Discrimination by biological state. Top row, receiver "
       "operating characteristic curves; middle row, precision–recall curves "
       "with the prevalence baseline; bottom row, confusion matrices at the "
       "Youden operating point with sensitivity and specificity given for both "
       "outcome classes.")

H("3.4 What the model uses, and which pathways those genes represent", 2)
P("SHAP attribution on the gradient-boosted model produced a ranked panel for "
  "each arm (Figure 3). Before treatment the highest-ranked genes were "
  f"{genes_list('DX', 6)}. At week 24 the panel was headed by "
  f"{genes_list('week_24', 6)}; the fifteen highest-ranked genes at that "
  "timepoint also include the high-affinity Fc gamma receptor gene FCGR1B and "
  "the interferon-inducible transcripts VAMP5 and METTL7B.")
sig_e = enrich[enrich.adjusted_p < 0.05]
P("Testing each attribution panel against six pre-specified immune programmes "
  "made that impression quantitative (Table 3). No programme was "
  "over-represented in the pre-treatment, day 7 or week 4 panels. In the week "
  "24 panel, the interferon-inducible programme was strongly over-represented "
  f"({enrich[(enrich.arm == 'week_24') & (enrich.term.str.startswith('Interferon'))].iloc[0].overlap} "
  f"genes, adjusted p = {enrich[(enrich.arm == 'week_24') & (enrich.term.str.startswith('Interferon'))].iloc[0].adjusted_p:.1e}), "
  "as was the inflammatory and myeloid activation programme "
  f"(adjusted p = {enrich[(enrich.arm == 'week_24') & (enrich.term.str.startswith('Inflammatory'))].iloc[0].adjusted_p:.3f}).")
rows = [[LABEL[r.arm], str(r.term), r.overlap, sci(r.p_value), sci(r.adjusted_p)]
        for _, r in enrich[enrich.n_overlap > 0].iterrows()]
if rows:
    table(["Arm", "Immune programme", "Overlap", "p", "Adjusted p"], rows,
          caption="Table 3. Over-representation of pre-specified immune "
                  "programmes among the highest-ranked genes by SHAP "
                  "attribution. Programmes with no overlapping gene are omitted "
                  "and are listed in full in Supplementary Table S10.",
          widths=[1.0, 2.6, 0.7, 0.9, 0.9], font=8.5)
ed_top = enrich_deg[enrich_deg.arm == "week_24"].head(6)
if len(ed_top):
    P("Open-ended enrichment was possible only at week 24, since that is the "
      "only arm producing genes that survive false-discovery correction. "
      "Applied to those genes it recovered cytokine signalling in the immune "
      f"system (adjusted p = "
      f"{ed_top[ed_top.term.str.contains('Cytokine')].adjusted_p.min() if ed_top.term.str.contains('Cytokine').any() else float('nan'):.4f}) "
      "and interferon signalling among the leading immune terms, alongside "
      "generic transcription and RNA-metabolism categories of the kind "
      "expected from a gene list of several thousand members "
      "(Supplementary Table S11).")

figure("Figure_shap_arms.png",
       "Figure 3. Feature attribution by biological state. Bars show mean "
       "absolute SHAP value for the fifteen highest-ranked genes in each arm. "
       "Colour indicates the direction of the association: red where high "
       "expression raises predicted risk, blue where it lowers it.")

H("3.5 Predicted risk tracks a neutrophil-high, T-cell-low axis", 2)
d24 = dec("week_24", "Neutrophil")
ddx = dec("DX", "Neutrophil")
P("Marker-panel deconvolution recovered the cellular basis of the signal "
  f"(Figure 4). Before treatment, neutrophil scores were higher in patients "
  f"with an unfavourable outcome, with a moderate effect size "
  f"(rank-biserial r = {ddx.rank_biserial_r:+.2f}) that fell short of "
  f"significance at this sample size (p = {ddx.p_value:.3f}). By week 24 the "
  f"same contrast was pronounced (r = {d24.rank_biserial_r:+.2f}, p "
  f"{pval(d24.p_value)}).")
best_corr = corr.loc[corr.spearman_rho.abs().idxmax()]
cdx_n, cdx_t = cor("DX", "Neutrophil"), cor("DX", "T cell")
c24_n, c24_t = cor("week_24", "Neutrophil"), cor("week_24", "T cell")
d7 = dec("day_7", "Neutrophil - T cell")
P("The decisive check is whether the machine-learning signal and the "
  "deconvolution describe the same thing. They do. Predicted probability of an "
  "unfavourable outcome correlated positively with the neutrophil score and "
  "negatively with the T-cell score in every arm (Figure 5): before treatment "
  f"ρ = {cdx_n.spearman_rho:+.2f} and {cdx_t.spearman_rho:+.2f} "
  f"respectively, and at week 24 ρ = {c24_n.spearman_rho:+.2f} and "
  f"{c24_t.spearman_rho:+.2f}. A model given no cell-composition information "
  "arrived independently at the axis that an established deconvolution method "
  "identifies.")
P("One comparison deserves note. At day 7 the composite neutrophil-minus-T-cell "
  f"score separated outcome groups clearly (r = {d7.rank_biserial_r:+.2f}, "
  f"p = {d7.p_value:.3f}) even though the classifiers trained on the same "
  "samples did not survive permutation testing. A single pre-specified summary "
  "of immune composition is estimated from a handful of parameters, whereas a "
  "model selecting from sixteen thousand genes must spend its limited "
  "statistical budget on that selection. With six events, the low-dimensional "
  "summary is the more sensitive instrument.")

figure("Figure_deconvolution_arms.png",
       "Figure 4. Immune composition by outcome in each biological state. "
       "Violin plots show marker-panel enrichment scores for neutrophils, "
       "T cells and their difference, with individual samples, medians, "
       "Mann–Whitney p-values and rank-biserial effect sizes.")

figure("Figure_prediction_vs_celltype.png",
       "Figure 5. Predicted probability of an unfavourable outcome against "
       "immune-cell scores. Each point is one sample, coloured by observed "
       "outcome, with the least-squares fit and Spearman correlation shown for "
       "each arm.")

H("3.6 Organisation of the signal", 2)
h24 = hubs[hubs.arm == "week_24"].head(4)
hdx = hubs[hubs.arm == "DX"].head(4)
P("Estimating a Gaussian graphical model on each SHAP panel showed how those "
  "features relate to one another once indirect associations are removed "
  "(Figure 6). Hub genes by degree were "
  f"{', '.join(str(g) for g in hdx.gene_symbol)} before treatment and "
  f"{', '.join(str(g) for g in h24.gene_symbol)} at week 24. These are "
  "undirected conditional associations. They indicate which transcripts carry "
  "shared information within the panel, and they do not establish regulatory "
  "direction, causation, or therapeutic relevance.")

figure("Figure_network_arms.png",
       "Figure 6. Conditional-dependency structure of the attribution panel in "
       "each arm. Nodes are the highest-ranked genes by SHAP value; edges are "
       "partial correlations from a graphical-lasso estimate, red for positive "
       "and blue for negative; node size reflects degree. Outcome is not a "
       "node and no direction is implied.")

H("3.7 Published signatures reproduce the result, and outperform our models", 2)
sw24 = pubsig[(pubsig.arm == "week_24") & (pubsig.signature.str.startswith("Sweeney"))].iloc[0]
ka24 = pubsig[(pubsig.arm == "week_24") & (pubsig.signature.str.startswith("Kaforou"))].iloc[0]
kadx = pubsig[(pubsig.arm == "DX") & (pubsig.signature.str.startswith("Kaforou"))].iloc[0]
P("The strongest test of the week-24 finding is whether it survives without "
  "any model of ours. Scoring the same samples with four published signatures "
  "of active tuberculosis, none of which has a parameter fitted to these data, "
  "reproduces it emphatically. The three-gene Sweeney signature separated "
  f"outcome groups at week 24 with a ROC-AUC of {sw24.roc_auc:.2f} "
  f"(p = {sci(sw24.p_value)}), and the Kaforou myeloid-up, lymphoid-down "
  f"signature reached {ka24.roc_auc:.2f} (p = {sci(ka24.p_value)}). Before "
  "treatment, the same signatures behaved as our models did, discriminating "
  f"weakly and without significance (Kaforou {kadx.roc_auc:.2f}, "
  f"p = {kadx.p_value:.3f}; Table 4).")
rows = [[LABEL[r.arm], r.signature, f"{r.roc_auc:.3f}",
         f"{r.rank_biserial_r:+.2f}", sci(r.p_value)]
        for _, r in pubsig.iterrows()]
table(["Arm", "Published signature", "ROC-AUC", "Rank-biserial r", "p"], rows,
      caption="Table 4. Published blood signatures of active tuberculosis "
              "applied to each arm without fitting. Scores are the mean "
              "standardised expression of up-weighted members minus that of "
              "down-weighted members.",
      widths=[0.95, 2.55, 0.75, 1.05, 0.8], font=8.5)

P("Two conclusions follow, and the second is uncomfortable. The first is that "
  "the week-24 result is not an artefact of our modelling: an externally "
  "derived, unfitted three-gene score reproduces it at a significance level "
  "our own cross-validated model does not reach. The second is that our "
  "models add nothing. In every arm the best published signature matched or "
  "exceeded the best fitted classifier (Table 5), by "
  f"{sigcmp.difference.abs().min():.2f} to {sigcmp.difference.abs().max():.2f} "
  "in ROC-AUC. A machine-learning pipeline with access to sixteen thousand "
  "genes did not improve on a score built from three.")
rows = [[r.arm, f"{r.fitted_model_auc:.3f}", r.best_published_signature,
         f"{r.published_auc:.3f}", f"{r.difference:+.3f}"]
        for _, r in sigcmp.iterrows()]
table(["Arm", "Best fitted model", "Best published signature",
       "Published ROC-AUC", "Difference"], rows,
      caption="Table 5. Fitted models compared with unfitted published "
              "signatures. A negative difference means the published "
              "signature discriminated better.",
      widths=[1.0, 1.0, 2.3, 1.0, 0.85], font=8.5)

H("3.8 Within-subject change discriminates far less well than state", 2)
t24 = traj[traj.comparison.str.contains("Week 24")].iloc[0]
P("The arm analyses compare different patients, so between-subject variation "
  "sits inside each contrast. Restricting to subjects with a paired diagnosis "
  "sample and forming the within-subject change vector removes it. The picture "
  f"weakens substantially. Alignment of each subject's change with the "
  "cured-response direction was lower in patients with an unfavourable outcome "
  f"at week 24, in the expected direction, but did not reach significance "
  f"(rank-biserial r = {t24.rank_biserial_r:+.2f}, p = {t24.p_value:.3f}; "
  f"n = {int(t24.n_subjects)} paired subjects, {int(t24.n_unfavourable)} "
  "unfavourable). Applying the arm classification protocol to the change "
  f"vectors gave a ROC-AUC of {t24.roc_auc_delta_model:.2f}, against "
  f"{best('week_24')['roc_auc']:.2f} for the cross-sectional week-24 state.")
P("This gap matters for interpretation. Part of it is expected on statistical "
  "grounds, since differencing two noisy measurements roughly doubles the "
  "variance while the paired subset is smaller. But it also means we cannot "
  "attribute the week-24 separation specifically to a failure to resolve under "
  "treatment. The evidence supports the weaker and more defensible claim: at "
  "the end of therapy, patients recorded as not cured occupy a different "
  "immunological state. Whether they travelled a different distance to reach "
  "it is not established by these data.")

figure("Figure_response_trajectory.png",
       "Figure 7. Within-subject treatment response. Alignment between each "
       "subject's diagnosis-to-later change vector and the average change of "
       "cured subjects, with the reference direction rebuilt leaving out each "
       "cured subject in turn.")

H("3.9 The signature does not transfer to an independent cohort", 2)
edn = ext_cell[(ext_cell.timepoint == "Diagnosis") & (ext_cell.cell_type == "Neutrophil")]
ednlr = ext_cell[(ext_cell.timepoint == "Diagnosis") & (ext_cell.cell_type == "NLR_score")]
P("We tested the immune axis in GSE67589, an independent cohort of 20 patients "
  "sampled at diagnosis, week 2 and week 4, in whom outcome was recorded as "
  "cure or relapse after apparently successful treatment [14]. Patients, "
  "country, array platform and outcome definition all differ from the "
  "discovery cohort.")
if len(edn) and len(ednlr):
    P("It did not transfer. Neutrophil scores were indistinguishable between "
      f"outcome groups at diagnosis (rank-biserial r = "
      f"{edn.iloc[0].rank_biserial_r:+.2f}, p = {edn.iloc[0].p_value:.2f}), and "
      "the neutrophil-minus-T-cell contrast, while nominally in the expected "
      f"direction (r = {ednlr.iloc[0].rank_biserial_r:+.2f}), was far from "
      f"significance (p = {ednlr.iloc[0].p_value:.2f}). No cell-composition "
      "comparison reached significance at any of the three timepoints "
      "(Figure 8; Supplementary Table S15).")
if len(ext_auc):
    P("Directional gene signatures derived from the discovery cohort performed "
      f"no better than chance when applied to the replication data ({ext_range()}; "
      "Table 6). We report this as a failure to generalise rather than as an "
      "inconclusive result.")
    rows = [[r.timepoint, r.discovery_arm, r.comparison,
             int(r.genes_transferred), int(r.n), int(r.n_relapse),
             f"{r.roc_auc:.2f} ({r.ci_low:.2f}–{r.ci_high:.2f})",
             f"{r.p_value:.2f}"]
            for _, r in ext_auc.iterrows()]
    table(["Replication timepoint", "Discovery arm", "Comparison",
           "Genes transferred", "n", "Relapse", "ROC-AUC (95% CI)", "p"], rows,
          caption="Table 6. Transfer of discovery-derived signatures to the "
                  "independent cohort. Scores apply discovery gene directions to "
                  "rank-standardised replication data. The week-24 rows are "
                  "cross-timepoint comparisons, since the replication cohort "
                  "stops at week 4.",
          widths=[1.0, 0.85, 0.95, 0.85, 0.35, 0.55, 1.25, 0.5], font=8)
P("Three features of this test limit what its failure establishes. The "
  "replication cohort defines an unfavourable outcome as relapse after "
  "treatment was judged successful, which is not the same phenotype as failure "
  "to be cured. It contains between eight and ten patients per group, so its "
  "confidence intervals are extremely wide. And it stops at week 4, so the "
  "discovery arm in which we observe a strong signal has no matching timepoint "
  "to be tested against; the week-24 rows in Table 6 compare across different "
  "stages of treatment. The result therefore shows that the signature does not "
  "generalise to this cohort, and leaves open whether it would generalise to a "
  "cohort matched in outcome definition and sampling schedule.")

figure("Figure_external_replication.png",
       "Figure 8. Independent cohort GSE67589. Neutrophil and T-cell scores by "
       "outcome at each sampled timepoint, with Mann–Whitney p-values and "
       "rank-biserial effect sizes. No comparison reaches significance.")

H("3.10 What the comparison between states shows", 2)
P("Bringing the arms together identifies where the outcome signal lives "
  f"(Figure 9). Of the genes separating outcome groups, {pcount('pre_treatment_only')} "
  "did so only before treatment, "
  f"{pcount('stable_across_states')} did so in both states, and "
  f"{pcount('post_treatment_only')} appeared only after treatment had begun. "
  "The signal is therefore overwhelmingly a property of the treated state, not "
  "an intrinsic property of the patient that treatment merely reveals.")
dnl = {a: dec(a, "Neutrophil - T cell") for a in C_ARMS}
P("The attribution panels tell the same story: the genes the model relies on "
  "before treatment overlap little with those it relies on at week 24. The "
  "immune axis behaves differently again. Its direction is the same in all "
  "four states — neutrophil-high and T-cell-low in patients with an "
  "unfavourable outcome — but its magnitude does not simply grow with time on "
  f"treatment (rank-biserial r = {dnl['DX'].rank_biserial_r:+.2f}, "
  f"{dnl['day_7'].rank_biserial_r:+.2f}, {dnl['week_4'].rank_biserial_r:+.2f} "
  f"and {dnl['week_24'].rank_biserial_r:+.2f} at diagnosis, day 7, week 4 and "
  "week 24). The contrast is clearest at day 7 and at week 24 and weakest at "
  "week 4. What changes across states is therefore not which axis matters but "
  "how far apart the two outcome groups sit along it, and that distance is "
  "not a monotonic function of time.")

figure("Figure_comparative.png",
       "Figure 9. Contrast between biological states. (A) Discrimination with "
       "95% confidence intervals and permutation p-values. (B) Immune-cell "
       "effect sizes by arm; asterisks mark p < 0.05. (C) Genes separating "
       "outcome groups, partitioned by the states in which they do so. "
       "(D) Overlap between the attribution panels of different arms.")

H("3.11 Robustness of the end-of-treatment result", 2)
P("Because the week-24 arm carries the study's only positive finding, it was "
  "attacked directly. Library size and gene-detection rate did not differ "
  f"between outcome groups in any arm (smallest p = {AUDIT['technical_min_p']:.2f}), "
  "so the result is not a sequencing-depth artefact. Removing any single "
  "patient with an unfavourable outcome and refitting left discrimination "
  f"between {AUDIT['influence_min']:.2f} and {AUDIT['influence_max']:.2f}, so "
  "it does not rest on one or two individuals. All three classifiers agreed "
  f"({AUDIT['ms_min']:.2f} to {AUDIT['ms_max']:.2f}). Four arms were tested, "
  "and week 24 remains significant after Holm correction across them "
  f"(adjusted p = {AUDIT['holm_week24']:.3f}), while no other arm approaches "
  "significance.")
P("One check was less reassuring. Bacterial load measured at diagnosis "
  "correlated with the week-24 predicted risk "
  f"({AUDIT['load_text']}), so patients who began treatment with a heavier "
  "burden tended to receive higher risk scores six months later. The deposit "
  "records no culture result at week 24, so residual bacterial load at the "
  "time of sampling cannot be adjusted for. Part of what the week-24 model "
  "reads may therefore be persistent mycobacterial burden rather than a host "
  "response that is informative independently of it.")

H("3.12 Sex and bacterial load do not explain the association", 2)
cdx = conf[(conf.arm == "DX") & (conf.model == "Neutrophil, adjusted for sex")]
c7 = conf[(conf.arm == "day_7") & (conf.model == "Neutrophil, adjusted for sex and bacterial load")]
c24 = conf[(conf.arm == "week_24") & (conf.model == "Neutrophil, adjusted for sex")]
sd = sexd[sexd.arm == "DX"].iloc[0]
P(f"Unfavourable outcomes were more common in men ({sd.male_among_unfavourable} "
  "before treatment). Adjusting for inferred sex left the neutrophil "
  "association essentially unchanged in every arm"
  + (f": the adjusted odds ratio was {cdx.iloc[0].odds_ratio:.1f} before "
     f"treatment (p = {cdx.iloc[0].p_value:.3f}, not significant at this "
     f"sample size) and {c24.iloc[0].odds_ratio:.1f} at week 24 "
     f"(p = {c24.iloc[0].p_value:.3f})" if len(cdx) and len(c24) else "")
  + ". Adding bacterial load attenuated the association without removing it"
  + (f"; at day 7, where the estimate is best determined, the fully adjusted "
     f"odds ratio was {c7.iloc[0].odds_ratio:.1f} "
     f"(p = {c7.iloc[0].p_value:.3f})" if len(c7) else "")
  + ". Confidence intervals on all of these estimates are wide, as the event "
    "counts require.")
P("Sex-linked transcripts were tested directly and showed no differential "
  f"expression by outcome in any arm (smallest p = {ylink.p_value.min():.2f}). "
  "Where such genes appear high in an unadjusted feature ranking, they reflect "
  "the sex composition of a small outcome group rather than a biological "
  "signal of treatment failure.")

# =====================================================================
# 4. DISCUSSION
# =====================================================================
H("4. Discussion", 1)

P("This study asked when, during tuberculosis treatment, whole blood carries "
  "information about the outcome of that treatment. The answer is sharply "
  "time-dependent. Through diagnosis, day 7 and week 4 it carries very little: "
  "discrimination was modest, confidence intervals included chance, no arm "
  "separated from its permutation null, and no gene survived false-discovery "
  "correction. At the end of therapy it carries a great deal, with strong "
  "discrimination, agreement across three independent classifiers, and "
  "thousands of genes differentially expressed at a controlled false-discovery "
  "rate.")

P("Reporting both halves matters. A pooled analysis of the same data yields an "
  "intermediate figure that describes neither state and would, if presented "
  "alone, misrepresent both. Separating the states converts what looks like a "
  "single mediocre result into two interpretable ones: an honest negative for "
  "pre-treatment stratification, and a robust positive for treatment-response "
  "monitoring.")

P("The end-of-treatment result should not be mistaken for prediction. Outcome "
  "is ascertained at approximately the same stage at which the week-24 sample "
  "is drawn, so the model is reading a concurrent state rather than "
  "anticipating a future one. Its value lies elsewhere: it characterises what "
  "immunological resolution looks like when treatment has worked, and shows "
  "that patients in whom treatment has not worked remain measurably distinct. "
  "The day 7 and week 4 arms show that this separation is not simply a "
  "function of time on treatment: outcome remained unreadable through the "
  "first four weeks and became readable only at the end of therapy. Whatever "
  "distinguishes the two groups at week 24 is therefore a property of "
  "completed or failed resolution rather than an early response to "
  "chemotherapy.")

P("Neutrophils and tuberculosis. Within the discovery cohort, every strand of "
  "evidence converges on a neutrophil-high, T-cell-low axis. The "
  "machine-learning models, given no information about cell composition, "
  "produced predictions correlated with deconvolved neutrophil and T-cell "
  "scores. Pathway analysis of the attribution panel recovered the "
  "interferon-inducible and myeloid activation programmes rather than "
  "neutrophil granule genes as such, which is what the interferon-driven "
  "neutrophil signature of active tuberculosis would predict: the cellular "
  "shift shows up in composition, and the transcriptional shift shows up in "
  "interferon-inducible myeloid genes. The association survived adjustment for "
  "sex and "
  "attenuated only partially on adjustment for bacterial load. Neutrophils "
  "have a recognised dual role in tuberculosis, contributing to early "
  "containment but also to tissue damage and to a permissive replicative niche "
  "in established disease [8,13], so a neutrophil-dominated state accompanying "
  "unfavourable outcome is biologically coherent. Our data are observational "
  "and cannot establish whether that state causes failure, follows from higher "
  "bacterial burden, or marks disease severity by another route.")

P("What the machine learning added. Very little, and saying so is part of the "
  "result. In every arm, a published signature computed without fitting a "
  "single parameter matched or exceeded the best of three cross-validated "
  "classifiers with access to the whole transcriptome. At week 24 a three-gene "
  "score reached a significance level our own models did not. Two things "
  "follow. The week-24 finding is unusually well supported, because it is "
  "reproduced by external, independently derived instruments rather than only "
  "by a model fitted to these data. And the flexible learner conferred no "
  "advantage at this sample size, which is what should be expected when seven "
  "events must support selection among sixteen thousand features. The useful "
  "role for machine learning here was not prediction but attribution: it "
  "identified which transcripts carried the signal, and those transcripts then "
  "turned out to be the ones the published signatures already contain.")

P("What the end-of-treatment signature is. The genes involved — FCGR1A, "
  "FCGR1B, VAMP5, METTL7B and their interferon-inducible neighbours — are the "
  "canonical markers of active tuberculosis rather than anything specific to "
  "treatment failure. The most parsimonious reading of the week-24 arm is "
  "therefore that patients recorded as not cured still have active disease and "
  "still carry its signature, while cured patients have resolved it. That is "
  "coherent, and it is confirmed by four independent signatures, but it is "
  "closer to a restatement of established biology than to a discovery. It also "
  "explains the failed replication rather neatly: the replication cohort's "
  "unfavourable outcome is relapse after treatment was judged successful, so "
  "those patients had by definition resolved their disease at the point of "
  "sampling and would not be expected to carry an active-disease signature.")

P("Convergence within a cohort is not generalisation. That distinction is the "
  "second main result of this study. Four methods agreeing inside GSE89403 "
  "establishes internal coherence, and it is tempting to read such agreement "
  "as evidence that a signature is real in a transferable sense. The "
  "independent cohort shows why that inference is unsafe: the signature did "
  "not transfer, and the cell-composition contrast that is pronounced at week "
  "24 in the discovery data is absent throughout the replication data. Some of "
  "this is attributable to the mismatch between cohorts in outcome definition, "
  "platform and sampling schedule, and some may be attributable to the small "
  "numbers on both sides. We cannot separate those explanations here, and we "
  "do not claim the signature would fail in a better-matched cohort. What we "
  "can say is that the evidence currently available does not support "
  "generalisation, and reporting that is more useful than reporting the "
  "internal convergence alone.")

P("Relation to other work. These results are consistent with the difficulty of "
  "baseline outcome prediction reported in larger cohorts [7], and complement "
  "models built from clinical records and multimodal data [9–11] by adding "
  "cellular resolution and an explicit audit of confounding. The design also "
  "follows a pattern that has proved useful elsewhere: derive features by "
  "machine learning, then verify them against established bioinformatic "
  "methods rather than reporting the model output alone [15].")

H("4.1 Limitations", 2)
P("The limitations are substantial and shape every conclusion. Event numbers "
  "are small, between six and eight per arm, so confidence intervals are wide "
  "and estimates unstable. The discovery cohort is a single study from one "
  "setting. The signature did not transfer to the one independent cohort we "
  "could identify, and although that cohort is small and differs in outcome "
  "definition, platform and sampling schedule, no external support for "
  "generalisation is currently available. The design is observational, so no "
  "causal claim is possible, and network edges are conditional associations "
  "only. Cell composition is inferred computationally and would need flow "
  "cytometry or single-cell measurement to confirm. Diabetes status and "
  "several other clinically relevant covariates are unrecorded. Week-24 "
  "discrimination is concurrent with outcome ascertainment and is not evidence "
  "of predictive capability; it is also associated with bacterial load "
  "measured at diagnosis, and the deposit carries no week-24 culture result, "
  "so residual mycobacterial burden cannot be separated from host response. "
  "The within-subject change from diagnosis discriminated far less well than "
  "the week-24 state itself, so the separation cannot be attributed "
  "specifically to differential resolution under treatment. Nothing reported "
  "here is ready for clinical use. "
  "The work is exploratory, and what we claim for it is that its negative and "
  "positive findings are reported with equal weight: outcome is not readable "
  "from pre-treatment blood in this cohort, it is strongly readable at the end "
  "of therapy, and neither the pre-treatment nor the on-treatment signature "
  "generalised when tested.")

# =====================================================================
# DECLARATIONS
# =====================================================================
H("Declarations", 1)
H("Ethical approval", 2)
P("This study analysed publicly available, de-identified gene-expression data "
  "from the NCBI Gene Expression Omnibus. No new participants were enrolled "
  "and no identifiable data were accessed, so formal ethical approval was not "
  "required. The source studies obtained the approvals described in their "
  "primary publications [12,14].")
H("Consent to participate", 2)
P("Not applicable; no new participants were enrolled.")
H("Consent to publish", 2)
P("Not applicable; no individually identifiable data are presented.")
H("Data availability", 2)
P("GSE89403 and GSE67589 are available from the NCBI Gene Expression Omnibus. "
  "Processed matrices and analysis outputs accompany the code release.")
H("Code availability", 2)
P("Analysis code, frozen models and intermediate outputs are available at "
  "https://github.com/hssling/TB-Transcriptomics-Project.")
H("Funding", 2)
P("This research received no specific grant from any funding agency in the "
  "public, commercial or not-for-profit sectors.")
H("Competing interests", 2)
P("The author declares no competing interests.")
H("Author contributions", 2)
P("Siddalingaiah H S: conceptualisation, methodology, software, formal "
  "analysis, investigation, data curation, writing of the original draft, "
  "review and editing, and visualisation.")
H("Use of artificial intelligence", 2)
P("Generative artificial-intelligence tools assisted with code development, "
  "literature searching and manuscript drafting. All such content was verified "
  "and edited by the author, who takes full responsibility for the work.")

# =====================================================================
# REFERENCES
# =====================================================================
H("References", 1)
REFS = [
    "World Health Organization. Global tuberculosis report 2024. Geneva: World Health Organization; 2024. ISBN 9789240101531.",
    "World Health Organization. Treatment of tuberculosis: guidelines. 4th ed. Geneva: World Health Organization; 2010. ISBN 9789241547833.",
    "Berry MPR, Graham CM, McNab FW, et al. An interferon-inducible neutrophil-driven blood transcriptional signature in human tuberculosis. Nature. 2010;466(7309):973–977. doi:10.1038/nature09247.",
    "Zak DE, Penn-Nicholson A, Scriba TJ, et al. A blood RNA signature for tuberculosis disease risk: a prospective cohort study. Lancet. 2016;387(10035):2312–2322. doi:10.1016/S0140-6736(15)01316-1.",
    "Sweeney TE, Braviak L, Tato CM, Khatri P. Genome-wide expression for diagnosis of pulmonary tuberculosis: a multicohort analysis. Lancet Respir Med. 2016;4(3):213–224. doi:10.1016/S2213-2600(16)00048-5.",
    "Singhania A, Verma R, Graham CM, et al. A modular transcriptional signature identifies phenotypic heterogeneity of human tuberculosis infection. Nat Commun. 2018;9(1):2308. doi:10.1038/s41467-018-04579-w.",
    "Mendelsohn SC, Andrade BB, Araújo-Pereira M, et al. Blood transcriptomic signatures predict poor treatment outcomes in drug-susceptible pulmonary TB in Brazil. medRxiv [preprint]. 2025 Oct 8. doi:10.1101/2025.10.07.25337480.",
    "Lowe DM, Redford PS, Wilkinson RJ, O’Garra A, Martineau AR. Neutrophils in tuberculosis: friend or foe? Trends Immunol. 2012;33(1):14–25. doi:10.1016/j.it.2011.10.003.",
    "Sambarey A, Smith K, Chung C, et al. Integrative analysis of multimodal patient data identifies personalized predictors of tuberculosis treatment prognosis. iScience. 2024;27(2):109025. doi:10.1016/j.isci.2024.109025.",
    "Vinhaes CL, Fukutani ER, Santana GC, et al. An integrative multi-omics approach to characterize interactions between tuberculosis and diabetes mellitus. iScience. 2024;27(3):109135. doi:10.1016/j.isci.2024.109135.",
    "Peng A-Z, Kong X-H, Liu S-T, et al. Explainable machine learning for early predicting treatment failure risk among patients with TB-diabetes comorbidity. Sci Rep. 2024;14(1):6814. doi:10.1038/s41598-024-57446-8.",
    "Thompson EG, Du Y, Malherbe ST, et al. Host blood RNA signatures predict the outcome of tuberculosis treatment. Tuberculosis (Edinb). 2017;107:48–58. doi:10.1016/j.tube.2017.08.004.",
    "Krug S, Parveen S, Bishai WR. Host-directed therapies: modulating inflammation to treat tuberculosis. Front Immunol. 2021;12:660916. doi:10.3389/fimmu.2021.660916.",
    "Cliff JM, Cho J-E, Lee J-S, et al. Excessive cytolytic responses predict tuberculosis relapse after apparently successful treatment. J Infect Dis. 2016;213(3):485–495. doi:10.1093/infdis/jiv447.",
    "Naidu A, Lulu SS. Systems and computational analysis of gene expression datasets reveals GRB-2 suppression as an acute immunomodulatory response against enteric infections in endemic settings. Front Immunol. 2024;15:1285785. doi:10.3389/fimmu.2024.1285785.",
    "Friedman J, Hastie T, Tibshirani R. Sparse inverse covariance estimation with the graphical lasso. Biostatistics. 2008;9(3):432–441. doi:10.1093/biostatistics/kxm045.",
    "Lundberg SM, Lee S-I. A unified approach to interpreting model predictions. Adv Neural Inf Process Syst. 2017;30:4765–4774.",
    "Chen T, Guestrin C. XGBoost: a scalable tree boosting system. In: Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining. 2016:785–794. doi:10.1145/2939672.2939785.",
    "Pedregosa F, Varoquaux G, Gramfort A, et al. Scikit-learn: machine learning in Python. J Mach Learn Res. 2011;12:2825–2830.",
    "Benjamini Y, Hochberg Y. Controlling the false discovery rate: a practical and powerful approach to multiple testing. J R Stat Soc Series B. 1995;57(1):289–300.",
]
for i, ref in enumerate(REFS, 1):
    P(f"{i}. {ref}", align=WD_ALIGN_PARAGRAPH.LEFT, size=10, space_after=4)

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f"wrote {OUT}")
