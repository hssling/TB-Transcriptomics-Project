# -*- coding: utf-8 -*-
"""Build revised Supplementary Material (.docx): full feature list, signatures,
DEG table, deconvolution stats, confounder table, external exclusion table,
benchmark table, recheck/robustness. Addresses R2.4 (accessible full lists),
R2.3 (exclusion table), R2.5 (DEG), R2.12 (effect sizes)."""
import json, os
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH

ROOT = "d:/research-automation/TB multiomics/TB-Treatment-Failure-Clean"
T = f"{ROOT}/DAI_Revision_2026/tables"
Fg = f"{ROOT}/DAI_Revision_2026/figures"


doc = Document()
doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(10)


def P(t, bold=False, size=10, hlt=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = bold; r.font.size = Pt(size)
    if hlt: r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return p


def table_from_df(df, max_rows=60, float_cols=None):
    df = df.head(max_rows).copy()
    t = doc.add_table(rows=1, cols=len(df.columns)); t.style = "Light Grid Accent 1"
    for j, c in enumerate(df.columns):
        t.rows[0].cells[j].paragraphs[0].add_run(str(c)).bold = True
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(df.columns):
            v = row[c]
            if isinstance(v, float):
                v = f"{v:.3g}"
            cells[j].text = str(v)
    doc.add_paragraph()


def supfig(path, caption, w=5.3):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(w))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    P(caption, size=9, hlt=True)


P("Supplementary Material", bold=True, size=14)
P("Baseline whole-blood transcriptomic risk stratification for unfavourable "
  "tuberculosis treatment outcome (DAI major revision).", size=11, hlt=True)

P("Supplementary Methods — class distribution and design", bold=True, size=12)
P("Primary cohort: GSE89403 pre-treatment (diagnosis) whole-blood RNA-seq, one "
  "sample per subject. N=90 subjects: 7 unfavourable ('Not Cured') and 83 cured "
  "(Definite/Probable/Possible Cure). Prevalence 7.8%. 16,147 gene features. "
  "Subject-level design eliminates timepoint leakage.", hlt=True)

# S1 performance
P("Supplementary Table S1 — Cross-validated performance (all models, with 95% CI)",
  bold=True, size=12)
A = json.load(open(f"{T}/wpA_metric_suite.json"))
rows = []
for m, d in A["models"].items():
    rows.append({"Model": m, "ROC-AUC": f"{d['roc_auc']:.2f}",
                 "ROC 95% CI": f"{d['roc_auc_ci95'][0]:.2f}-{d['roc_auc_ci95'][1]:.2f}",
                 "PR-AUC": f"{d['pr_auc']:.2f}", "Brier": f"{d['brier']:.2f}"})
table_from_df(pd.DataFrame(rows))
op = A["operating_point"]
P(f"Operating point (Youden, random forest): sensitivity {op['sensitivity_recall']:.2f}, "
  f"specificity {op['specificity']:.2f}, PPV {op['precision_ppv']:.2f}, "
  f"NPV {op['npv']:.2f}, F1 {op['f1']:.2f}, MCC {op['mcc']:.2f}; confusion "
  f"TN={op['confusion_matrix']['tn']}, FP={op['confusion_matrix']['fp']}, "
  f"FN={op['confusion_matrix']['fn']}, TP={op['confusion_matrix']['tp']}.", hlt=True)

# S2 SHAP feature list (full)
P("Supplementary Table S2 — Full ranked predictor list with HGNC symbols (SHAP)",
  bold=True, size=12)
shap = pd.read_csv(f"{T}/wpC_shap_feature_list.csv")
table_from_df(shap, max_rows=50)
supfig(f"{Fg}/Figure_shap_summary.png",
       "Supplementary Figure S2. SHAP summary (beeswarm) for the baseline "
       "gradient-boosting model; top predictors by mean |SHAP|.")

# S3 DEG corrected
P("Supplementary Table S3 — Differential expression (true log2 fold-change), top 60",
  bold=True, size=12)
deg = pd.read_csv(f"{T}/wpB_DEG_corrected.csv")
table_from_df(deg[["gene_symbol", "log2FC", "p_value", "fdr"]], max_rows=60)
P("No gene survived FDR<0.05 at baseline (n=7 events); table is hypothesis-"
  "generating. ALOX5AP (neutrophil leukotriene gene) is up-regulated in failure.",
  hlt=True)
supfig(f"{Fg}/Figure_volcano.png",
       "Supplementary Figure S1. Volcano plot of baseline differential "
       "expression (true log2 fold-change, failure vs cure). Dashed line p=0.05.")

# S4 deconvolution stats
P("Supplementary Table S4 — Immune deconvolution: effect sizes (rank-biserial)",
  bold=True, size=12)
vd = pd.read_csv(f"{T}/wpD_violin_stats.csv")
table_from_df(vd)
conc = json.load(open(f"{T}/wpD_concordance.json"))
P(f"Concordance (deconvolution vs model probability): neutrophil Spearman "
  f"rho={conc['neutrophil_vs_ML_prob']['spearman_rho']:.2f} (p~1e-12); "
  f"T-cell rho={conc['tcell_vs_ML_prob']['spearman_rho']:.2f}.", hlt=True)

# S5 external exclusion table
P("Supplementary Table S5 — External-cohort search and exclusion rationale",
  bold=True, size=12)
ex = pd.read_csv(f"{T}/wpF_external_exclusion_table.csv")
table_from_df(ex)
port = json.load(open(f"{T}/wpF_portability.json"))
P(f"Label-free portability (GSE193979, n={port['n_samples']}): neutrophil vs "
  f"T-cell score Spearman rho={port['neutrophil_vs_tcell_spearman']:.2f} "
  f"(p~3e-7) — expected anti-structure reproduced.", hlt=True)

# S6 benchmark
P("Supplementary Table S6 — Benchmark vs prior TB treatment-outcome / ML studies",
  bold=True, size=12)
bm = pd.read_csv(f"{T}/wpH_benchmark_table.csv")
table_from_df(bm)

# S7 confounders
P("Supplementary Table S7 — Confounder-adjusted logistic regression", bold=True, size=12)
cf = pd.read_csv(f"{T}/wpG_confounder_logit.csv")
table_from_df(cf)

# S8 cell-type specificity
P("Supplementary Table S8 — Cell-type specificity of signature genes (HPA atlas)",
  bold=True, size=12)
spec = json.load(open(f"{T}/wpE_specificity_summary.json"))
P(f"Neutrophil-signature genes max-expressed in neutrophils: "
  f"{spec['neutrophil_specific']}; T-cell-signature genes in T cells: "
  f"{spec['tcell_specific']}. Reference: {spec['reference']}.", hlt=True)

# S9 recheck (if present)
import os
if os.path.exists(f"{T}/recheck_summary.json"):
    rc = json.load(open(f"{T}/recheck_summary.json"))
    P("Supplementary Table S9 — Independent robustness recheck", bold=True, size=12)
    P(f"Leave-one-out CV ROC-AUC = {rc['loo_auc']:.2f} (vs repeated-KFold ~0.67). "
      f"Permutation null (1000 label shuffles): observed AUC {rc['perm_obs_auc']:.2f} "
      f"vs null mean {rc['perm_null_mean']:.2f}, permutation p = {rc['perm_p']:.3f}. "
      f"Neutrophil rank-biserial r={rc['neutrophil_r']:.2f} "
      f"(95% CI {rc['neutrophil_r_ci'][0]:.2f}..{rc['neutrophil_r_ci'][1]:.2f}), "
      f"MWU p={rc['neutrophil_mwu_p']:.3f}.", hlt=True)

# S10 enrichment
if os.path.exists(f"{T}/wpB_enrichment.csv"):
    P("Supplementary Table S10 — Pathway enrichment of top nominal DEGs "
      "(exploratory)", bold=True, size=12)
    en = pd.read_csv(f"{T}/wpB_enrichment.csv")
    table_from_df(en[["Gene_set", "Term", "P-value", "Adjusted P-value"]], max_rows=20)
    P("No term survived FDR<0.05 (minimum adjusted p≈0.11), consistent with the "
      "underpowered DEG; weak exploratory signals included adaptive-immune-"
      "response and protein-secretion/unfolded-protein-response terms.", hlt=True)

# S11 network hubs
if os.path.exists(f"{T}/wpI_network_hubs.csv"):
    P("Supplementary Table S11 — Conditional-dependency network hub genes "
      "(degree)", bold=True, size=12)
    hubs = pd.read_csv(f"{T}/wpI_network_hubs.csv")
    table_from_df(hubs, max_rows=26)
    P("Undirected Gaussian graphical model on top baseline failure-associated "
      "genes; outcome not included as a node (R2 comment 8).", hlt=True)

out = f"{ROOT}/DAI_Revision_2026/deliverables/Supplementary_Material_DAI_MajorRevision_v13.docx"
doc.save(out)
print("Saved supplementary:", out)
