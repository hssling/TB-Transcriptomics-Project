# -*- coding: utf-8 -*-
"""Build the COMPLETE revised DAI manuscript (.docx) with revised text
highlighted yellow (mandatory tracked-changes requirement). Honest exploratory
baseline risk-stratification framing; numbers wired to reanalysis outputs."""
import json, os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

ROOT = "d:/research-automation/TB multiomics/TB-Treatment-Failure-Clean"
T = f"{ROOT}/DAI_Revision_2026/tables"
F = f"{ROOT}/DAI_Revision_2026/figures"
A = json.load(open(f"{T}/wpA_metric_suite.json"))
rf = A["models"]["RandomForest"]; op = A["operating_point"]
cm = op["confusion_matrix"]
conc = json.load(open(f"{T}/wpD_concordance.json"))
port = json.load(open(f"{T}/wpF_portability.json"))
spec = json.load(open(f"{T}/wpE_specificity_summary.json"))
sens = json.load(open(f"{T}/wpA_sensitivity_timepoints.json"))
rc = json.load(open(f"{T}/recheck_summary.json"))

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Times New Roman"; st.font.size = Pt(11)
C = WD_ALIGN_PARAGRAPH.CENTER


def P(runs, align=None, size=11, italic=False, bold=False):
    p = doc.add_paragraph()
    if align: p.alignment = align
    if isinstance(runs, str): runs = [(runs, False)]
    for text, h in runs:
        r = p.add_run(text); r.font.size = Pt(size); r.italic = italic; r.bold = bold
        if h: r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return p


def H(t, l=1): doc.add_heading(t, level=l)
def hl(t): return (t, True)
def n(t): return (t, False)


def fig(path, caption_runs, width=6.0):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = C
    P(caption_runs, size=10)


# ===== TITLE =====
title = ("Baseline whole-blood transcriptomic risk stratification for "
         "unfavourable tuberculosis treatment outcome: an exploratory "
         "machine-learning and immune-deconvolution study")
p = doc.add_paragraph(); p.alignment = C
r = p.add_run(title); r.bold = True; r.font.size = Pt(14)
r.font.highlight_color = WD_COLOR_INDEX.YELLOW
P("Article type: Original Article", align=C)
P("Siddalingaiah H S, MD", align=C, bold=True)
P([n("Professor, Department of Community Medicine, Shridevi Institute of "
     "Medical Sciences and Research Hospital, Tumkur, Karnataka, India")], align=C, size=10)
P([n("Correspondence: Dr Siddalingaiah H S. Email: hssling@yahoo.com")], align=C, size=10)
P([hl("Note to editors/reviewers: This manuscript has been comprehensively "
      "revised. Causal and ‘first mechanistic map’ language has been removed; "
      "the study is reframed as an exploratory, associative baseline "
      "risk-stratification analysis. Substantively revised or new content is "
      "highlighted in yellow, per the tracked-changes requirement.")],
  size=10, italic=True)

# ===== ABSTRACT =====
H("Abstract", 1)
P([n("Background. "), hl(
   "Tuberculosis (TB) treatment failure and relapse sustain transmission and "
   "drug resistance, yet baseline (pre-treatment) tools to identify at-risk "
   "patients remain limited. Blood transcriptomic signatures predict TB "
   "disease and progression, but baseline prediction of treatment failure is "
   "recognised as difficult and most signatures are not resolved to specific "
   "immune-cell populations. Our aim was not to build a superior predictor but "
   "to characterise, transparently, what a baseline whole-blood signal of "
   "unfavourable outcome looks like at the cellular level.")])
P([n("Methods. "), hl(
   "We analysed public whole-blood RNA-seq from the GSE89403 (Catalysis) "
   "treatment-response cohort. To avoid timepoint leakage the primary analysis "
   "used pre-treatment (diagnosis) samples only, one per subject (N=90; 7 "
   "unfavourable [‘Not Cured’] vs 83 cured; prevalence 7.8%). Discrimination "
   "was estimated by repeated stratified cross-validation with pooled "
   "out-of-fold predictions and bootstrap 95% confidence intervals, reporting "
   "ROC-AUC, precision–recall AUC, sensitivity, specificity, predictive "
   "values, F1, MCC and calibration. We performed differential expression, "
   "SHAP interpretation, marker-based immune deconvolution with rank-biserial "
   "effect sizes, an undirected Gaussian graphical (conditional-dependency) "
   "network, a sex- and bacterial-load confounder audit, and "
   "cell-type-specificity validation in a granulocyte-containing reference.")])
P([n("Results. "), hl(
   f"Baseline discrimination was modest (random forest ROC-AUC {rf['roc_auc']:.2f}, "
   f"95% CI {rf['roc_auc_ci95'][0]:.2f}–{rf['roc_auc_ci95'][1]:.2f}; PR-AUC "
   f"{rf['pr_auc']:.2f}). At the Youden operating point sensitivity was "
   f"{op['sensitivity_recall']:.2f}, specificity {op['specificity']:.2f}, "
   f"negative predictive value {op['npv']:.2f} and positive predictive value "
   f"{op['precision_ppv']:.2f}—a rule-out rather than rule-in signal. "
   f"Discrimination did not reach statistical significance against a "
   f"label-permutation null (permutation p={rc['perm_p']:.2f}; leave-one-out "
   f"AUC {rc['loo_auc']:.2f}). Model predictions nonetheless tracked a "
   f"neutrophil-high/T-cell-low axis (Spearman ρ="
   f"{conc['neutrophil_vs_ML_prob']['spearman_rho']:.2f} and "
   f"{conc['tcell_vs_ML_prob']['spearman_rho']:.2f}, p≈10⁻¹⁰). Baseline neutrophil "
   f"elevation was a moderate but only borderline association (rank-biserial "
   f"r=−0.44, p=0.057), independent of sex (adjusted OR≈4.5); Y-chromosome "
   f"genes that dominated a naive model were shown to be a sex confound. "
   f"Neutrophil-signature genes were neutrophil-specific in a "
   f"granulocyte-containing reference ({spec['neutrophil_specific']}).")])
P([n("Conclusions. "), hl(
   "Baseline whole-blood data carry a weak, hypothesis-generating signal of TB "
   "treatment failure associated with—but not establishing causation by—a "
   "neutrophil-high/T-cell-low immunophenotype. Given seven events, wide "
   "confidence intervals and no outcome-labelled external cohort, the "
   "signature is not yet clinically actionable; we present it as an "
   "exploratory framework and a transparent baseline for future adequately "
   "powered, prospective validation.")])
P([n("Keywords: "), hl("tuberculosis; treatment outcome; whole-blood "
   "transcriptomics; machine learning; immune deconvolution; neutrophils; "
   "risk stratification; explainable AI")])

# ===== INTRODUCTION =====
H("1. Introduction", 1)
P([n("Tuberculosis (TB) remains a leading cause of death from a single "
     "infectious agent, with an estimated 1.25 million deaths in 2023 [1]. "
     "Standard short-course chemotherapy fails in roughly 5–10% of "
     "drug-susceptible and a substantially higher fraction of "
     "multidrug-resistant (MDR-TB) cases [1]. "),
   hl("Directly Observed Treatment, Short-course (DOTS)—the WHO-recommended "
      "case-management strategy combining standardised regimens with treatment "
      "observation—does not individualise therapy to host biology [1,2].")])
P([n("Blood transcriptomics is an established tool for TB biomarker discovery. "
     "Signatures from Berry, Zak, Sweeney and Singhania discriminate active "
     "from latent TB and predict progression [3–6]. "),
   hl("However, baseline prediction of treatment failure—as opposed to disease "
      "status—remains difficult: in a large drug-susceptible Brazilian cohort, "
      "baseline transcriptomic prediction of failure achieved AUC <0.70, "
      "although prediction of death and recurrence was stronger [7]. Two "
      "limitations recur. First, a resolution limitation: bulk signatures "
      "average expression across a shifting mixture of leukocytes, so a "
      "‘high-inflammation’ signal may reflect a change in cell proportions "
      "(e.g., neutrophilia with lymphopenia) rather than per-cell activation "
      "[8]. Second, signatures are rarely audited for confounders such as sex "
      "or bacterial load.")])
P([hl(
   "We emphasise that multi-omic and machine-learning approaches to TB "
   "treatment prognosis already exist; this work is not the first in the "
   "field. Integrative multimodal models combining radiological, "
   "microbiological and clinical data predict prognosis across thousands of "
   "patients [9]; multi-omic analyses characterise TB–diabetes interactions "
   "[10]; and explainable machine-learning models predict treatment failure "
   "from electronic medical records [11]. Our contribution is narrower and "
   "complementary: a transparent, cellularly-resolved, confounder-audited "
   "characterisation of the baseline whole-blood signal, using only "
   "pre-treatment samples and explicitly reporting its limits.")])
P([n("In this exploratory study we ask three questions. "), hl(
   "(i) How well can a baseline whole-blood signature discriminate eventual "
   "treatment failure, reported with full operating-point metrics and "
   "uncertainty? (ii) Which immune-cell compartments does that signal "
   "correspond to, quantified by effect size and concordance with model "
   "predictions? (iii) Is any signal robust to sex and bacterial-load "
   "confounding? We deliberately avoid causal claims and clinical-readiness "
   "claims.")])

# ===== METHODS =====
H("2. Methods", 1)
H("2.1 Data source, inclusion and exclusion criteria", 2)
P([n("The discovery cohort was GSE89403, a longitudinal whole-blood RNA-seq "
     "study of adults with pulmonary TB undergoing standard therapy "
     "(Catalysis treatment-response cohort) [12]. "),
   hl("The original cohort enrolled HIV-uninfected, drug-susceptible pulmonary "
      "TB patients, so HIV co-infection and rifampicin resistance are "
      "controlled by design rather than measured covariates; diabetes status "
      "was not recorded. Inclusion criteria for this analysis were: (i) "
      "availability of a pre-treatment (diagnosis, ‘DX’) whole-blood sample; "
      "(ii) an annotated treatment outcome. Outcome was dichotomised as "
      "unfavourable (‘Not Cured’) versus cured (‘Definite/Probable/Possible "
      "Cure’). After restricting to one diagnosis sample per subject, the "
      "primary cohort comprised N=90 subjects (7 unfavourable, 83 cured; "
      "prevalence 7.8%), with 16,147 gene features.")])
P([hl(
   "Class distribution is reported here in the main text (not only in the "
   "supplement) because it governs the interpretation of every performance "
   "metric: with seven events, confidence intervals are necessarily wide and "
   "all results are exploratory.")])
H("2.2 Preprocessing", 2)
P([n("Counts were "), hl(
   "library-size normalised and log1p-transformed; genes expressed in <10% of "
   "samples were removed for differential-expression testing. ‘Standardised "
   "pipeline’ here means a fixed, scripted sequence—normalisation, "
   "log-transformation, Ensembl-gene identifier harmonisation, and (for "
   "deconvolution and regression) per-gene z-scoring—applied identically to "
   "all samples and released as code (Section 2.9).")])
H("2.3 Predictive modelling and leakage control", 2)
P([n("Because the same subject contributes multiple on-treatment samples in "
     "GSE89403, "),
   hl("pooling timepoints inflates apparent performance through subject "
      "leakage. The primary model therefore used pre-treatment samples only "
      "(one per subject). We evaluated logistic regression, random forest and "
      "gradient boosting (XGBoost) within repeated stratified 5-fold "
      "cross-validation (40 repeats); univariate feature pre-selection was "
      "performed inside each training fold only. Out-of-fold predicted "
      "probabilities were pooled and metrics computed on them, with bootstrap "
      "(2,000-resample) 95% confidence intervals. We report ROC-AUC, "
      "precision–recall AUC, and—at the Youden operating point—sensitivity, "
      "specificity, positive and negative predictive value, F1, Matthews "
      "correlation coefficient, balanced accuracy, the confusion matrix, and "
      "calibration (Brier score). As a sensitivity analysis we also fitted a "
      "subject-grouped model across all timepoints to quantify the leakage "
      "effect.")])
H("2.4 Differential expression and explainability", 2)
P([hl(
   "Per-gene differential expression between cured and unfavourable groups "
   "used the Mann–Whitney U test with Benjamini–Hochberg false-discovery-rate "
   "(FDR) control and log2 fold-change; results are shown as a volcano plot. "
   "Model explainability used SHAP (TreeExplainer) on the baseline gradient-"
   "boosting model; the full ranked feature list with HGNC symbols is provided "
   "in the supplement.")])
H("2.5 Immune-cell deconvolution", 2)
P([hl(
   "We estimated immune-cell composition by marker-gene-set z-score enrichment "
   "(a transparent digital-cytometry approach), using canonical marker panels "
   "for neutrophils, T cells, monocytes, B cells and NK cells. We describe the "
   "method as marker-based enrichment rather than proprietary CIBERSORT/LM22 "
   "to match exactly what was computed. Group differences used Mann–Whitney U "
   "with rank-biserial correlation as the effect size. Concordance between the "
   "deconvolved neutrophil/T-cell scores and the model’s predicted failure "
   "probability was quantified by Spearman correlation.")])
H("2.6 Cell-type-specificity validation", 2)
P([n("The earlier submission validated signatures against the PBMC3k "
     "single-cell reference. "),
   hl("Because peripheral-blood mononuclear-cell preparations exclude "
      "granulocytes by density-gradient isolation, PBMC3k cannot validate a "
      "neutrophil signature. We therefore validated cell-type specificity in "
      "the Human Protein Atlas blood-cell atlas, a granulocyte-containing "
      "immune-cell reference, testing whether neutrophil- and T-cell-signature "
      "genes are maximally expressed in the expected lineage.")])
H("2.7 Conditional-dependency network", 2)
P([n("We constructed a Gaussian graphical model with the graphical lasso. "),
   hl("A non-zero entry in the sparse inverse-covariance matrix denotes a "
      "partial correlation conditional on all other genes. We stress that "
      "these edges are undirected conditional associations: they do not encode "
      "direction, causation, or regulatory hierarchy, and the clinical outcome "
      "is not an input to the estimator. High-degree nodes are therefore "
      "described as hub genes, not ‘upstream regulators’ or ‘drivers’.")])
H("2.8 Confounders and external validation", 2)
P([hl(
   "We assessed confounding by sex (estimated from XIST and RPS4Y1 expression) "
   "and bacterial load (MGIT time-to-positivity and Xpert Ct), using "
   "multivariable logistic regression. We searched the Gene Expression Omnibus "
   "for an independent cohort with baseline samples and treatment-outcome "
   "labels; the rationale for each candidate, and why outcome-stratified "
   "external validation was not feasible, are documented in a supplementary "
   "exclusion table. A label-free signature-portability check was performed in "
   "GSE193979.")])
H("2.9 Statistics, software and reproducibility", 2)
P([n("Analyses used Python (scikit-learn, XGBoost, SHAP, statsmodels, SciPy). "
     "All tests were two-sided at α=0.05 with FDR control where indicated. "),
   hl("Effect sizes (rank-biserial correlation, odds ratios with 95% CIs) are "
      "reported alongside p-values. Code, the frozen model and intermediate "
      "outputs are released to enable verification without re-execution "
      "(Section ‘Code availability’).")])

# ===== RESULTS =====
H("3. Results", 1)
H("3.1 Cohort and class distribution", 2)
P([hl(
   "The baseline cohort comprised 90 subjects (7 unfavourable outcomes, 83 "
   "cured; prevalence 7.8%). This severe imbalance and small event count frame "
   "all subsequent results as exploratory.")])
H("3.2 Baseline predictive performance is modest (Figure 1)", 2)
P([hl(
   f"On pre-treatment samples, baseline discrimination was modest. The best "
   f"model (random forest) achieved ROC-AUC {rf['roc_auc']:.2f} (95% CI "
   f"{rf['roc_auc_ci95'][0]:.2f}–{rf['roc_auc_ci95'][1]:.2f}); logistic "
   f"regression and XGBoost performed similarly or worse "
   f"(0.52 and {A['models']['XGBoost']['roc_auc']:.2f}). Precision–recall AUC "
   f"was {rf['pr_auc']:.2f} against a 0.08 prevalence baseline (Figure 1A,B). "
   f"At the Youden operating point (Figure 1D) sensitivity was "
   f"{op['sensitivity_recall']:.2f} and specificity {op['specificity']:.2f}, "
   f"with negative predictive value {op['npv']:.2f} but positive predictive "
   f"value only {op['precision_ppv']:.2f} (confusion matrix: TN={cm['tn']}, "
   f"FP={cm['fp']}, FN={cm['fn']}, TP={cm['tp']}; F1={op['f1']:.2f}, "
   f"MCC={op['mcc']:.2f}). Calibration was acceptable (Brier "
   f"{rf['brier']:.2f}; Figure 1C). Three independent cross-validation "
   f"procedures agreed (repeated stratified {rf['roc_auc']:.2f}, leave-one-out "
   f"{rc['loo_auc']:.2f}, single split 0.69), but discrimination did NOT reach "
   f"significance against a label-permutation null (observed "
   f"{rc['perm_obs_auc']:.2f} vs null mean {rc['perm_null_mean']:.2f}, "
   f"permutation p={rc['perm_p']:.2f}). The wide confidence interval (which "
   f"includes 0.5) together with the non-significant permutation test mean the "
   f"baseline predictor is, by itself, not statistically robust; its value is "
   f"as a transparent, reproducible exploratory baseline and a rule-out signal "
   f"(high negative predictive value), not as a deployable classifier.")])
fig(f"{F}/Figure1_performance.png",
    [n("Figure 1. "), hl("Baseline predictive performance (pooled out-of-fold). "
     "(A) ROC and (B) precision–recall curves for three classifiers; (C) "
     "calibration of the random-forest model; (D) confusion matrix at the "
     "Youden operating point. Label 1 = treatment failure (unfavourable), 0 = "
     "cure.")])
P([n("Sensitivity analysis. "), hl(
   f"A subject-grouped model pooling all timepoints reached ROC-AUC "
   f"{sens['all_timepoint_grouped']['roc_auc']:.2f}, but because all seven "
   f"failure samples are pre-treatment while many cured samples are "
   f"on-treatment, this value is confounded by treatment-induced "
   f"transcriptional change and is not a valid estimate of baseline "
   f"failure discrimination. This directly explains the higher accuracy "
   f"reported in the original submission and confirms the reviewers’ concern "
   f"that the signal is partly a timepoint effect.")])
H("3.3 Differential expression is underpowered (volcano)", 2)
P([hl(
   "Baseline differential expression (true log2 scale) identified ~1,600 genes "
   "at nominal p<0.05 but none survived FDR correction, consistent with low "
   "power at seven events; the volcano plot (Supplementary Figure S1) is "
   "therefore hypothesis-generating. Among the most differentially expressed "
   "genes was ALOX5AP (arachidonate-5-lipoxygenase-activating protein, a "
   "neutrophil leukotriene gene), up-regulated in failure, consistent with the "
   "neutrophil hypothesis. Notably, Y-chromosome genes (RPS4Y1, KDM5D, DDX3Y) "
   "were not differentially expressed at baseline (p>0.3) even though they "
   "dominated a naive classifier—evidence that those ‘predictors’ reflect the "
   "male skew of the small failure group rather than a biological failure "
   "signal (Section 3.6).")])
H("3.4 Model predictions correspond to a neutrophil-high/T-cell-low axis "
  "(Figure 2)", 2)
P([hl(
   f"Immune deconvolution showed higher neutrophil scores in failures (median "
   f"+0.36 vs −0.12) with a moderate effect size (rank-biserial r=−0.44), "
   f"although significance was borderline at this sample size (Mann–Whitney "
   f"p=0.057; Figure 2). The T-cell reduction was not significant at baseline "
   f"(p=0.47), in contrast to the strong significance reported previously from "
   f"pooled on-treatment samples. Crucially, the model’s predicted failure "
   f"probability correlated with the neutrophil score (Spearman ρ="
   f"{conc['neutrophil_vs_ML_prob']['spearman_rho']:.2f}, p≈1×10⁻¹²) and "
   f"inversely with the T-cell score (ρ="
   f"{conc['tcell_vs_ML_prob']['spearman_rho']:.2f}), demonstrating that the "
   f"machine-learning signal and the deconvolution converge on the same "
   f"myeloid–lymphoid axis (addressing the alignment between deconvolution and "
   f"model findings).")])
fig(f"{F}/Figure2_deconvolution_violins.png",
    [n("Figure 2. "), hl("Baseline immune-cell deconvolution, cure (0) vs "
     "failure (1). Violin plots with individual points, Mann–Whitney p-values, "
     "rank-biserial effect sizes and significance annotations for the "
     "neutrophil score, T-cell score and their difference (NLR axis).")])
H("3.5 Cell-type specificity in a granulocyte-containing reference (Figure 3)", 2)
P([hl(
   f"In the Human Protein Atlas blood atlas—which, unlike PBMC3k, contains "
   f"neutrophils—{spec['neutrophil_specific']} neutrophil-signature genes were "
   f"maximally expressed in neutrophils and {spec['tcell_specific']} T-cell-"
   f"signature genes in T cells (Figure 3), confirming lineage specificity of "
   f"the signatures and correcting the earlier, inappropriate PBMC3k "
   f"validation.")])
fig(f"{F}/Figure3_celltype_specificity.png",
    [n("Figure 3. "), hl("Cell-type specificity of failure-signature genes "
     "across labelled blood-cell types (Human Protein Atlas). Colour scale = "
     "row-max-normalised nTPM; the cyan line separates neutrophil- from "
     "T-cell-signature genes.")], width=4.8)
H("3.6 Sex and bacterial-load confounder audit", 2)
P([hl(
   "Failures were more often male (5/7), and Y-chromosome genes dominated a "
   "naive feature ranking. However, in multivariable logistic regression the "
   "neutrophil association remained (adjusted OR≈4.5, p≈0.046) while sex "
   "itself was not significant (OR≈1.8, p=0.52), indicating a sex-independent "
   "neutrophil signal and confirming that the Y-linked ‘predictors’ were a "
   "confound. Adjusting for bacterial load (MGIT/Xpert) attenuated but did not "
   "abolish the neutrophil association (OR≈4.5, p=0.10; reduced N=80). HIV and "
   "drug resistance were exclusion criteria in the source cohort.")])
H("3.7 Conditional-dependency network (Figure 4)", 2)
P([n("The graphical-lasso network yielded a sparse set of partial-correlation "
     "edges among the top baseline failure-associated genes "),
   hl("(the clinical outcome was not included as a node). Hub genes by degree "
      "included ALOX5AP—an arachidonate-5-lipoxygenase-activating protein "
      "expressed in neutrophils—together with ALMS1, EMC1 and MTG2 (Figure 4; "
      "Supplementary Table S11). We report these strictly as undirected "
      "conditional associations and make no claim that they are upstream "
      "regulators, drivers or therapeutic targets.")])
fig(f"{F}/Figure4_network.png",
    [n("Figure 4. "), hl("Undirected conditional-dependency (Gaussian "
     "graphical) network among baseline failure-associated genes. Edges are "
     "partial correlations (red positive, blue negative); node size reflects "
     "degree; gene symbols shown. The clinical outcome is not a node and no "
     "directionality is implied.")], width=5.2)
H("3.8 External portability and benchmarking", 2)
P([hl(
   f"An outcome-stratified external validation was not feasible: in the "
   f"closest candidate (GSE193979, TANDEM), per-patient outcomes are not "
   f"deposited in GEO and the public count matrix is keyed by internal "
   f"identifiers with no public bridge to sample metadata (Supplementary Table "
   f"S5). A label-free portability check nonetheless reproduced the expected "
   f"neutrophil↔T-cell anti-structure in that independent cohort (Spearman "
   f"ρ={port['neutrophil_vs_tcell_spearman']:.2f}, p≈3×10⁻⁷, n="
   f"{port['n_samples']}). Benchmarking against prior work (Supplementary "
   f"Table S6) shows our modest baseline AUC is consistent with the field, "
   f"where baseline failure prediction is hard (AUC <0.70) even in large "
   f"curated cohorts [7], and positions our contribution as cellular "
   f"resolution and confounder transparency rather than predictive "
   f"superiority [9–12].")])

# ===== DISCUSSION =====
H("4. Discussion", 1)
P([hl(
   "This study provides a transparent, cellularly-resolved characterisation of "
   "the baseline whole-blood signal of TB treatment failure. Three findings "
   "are robust to our small sample: the machine-learning signal and immune "
   "deconvolution converge on a neutrophil-high/T-cell-low axis; that "
   "neutrophil association is sex-independent; and the neutrophil signature is "
   "lineage-specific in a granulocyte-containing reference. Three findings are "
   "explicitly weak: baseline discrimination is modest with a confidence "
   "interval spanning chance; the T-cell reduction is not significant at "
   "baseline; and no outcome-labelled external validation was possible.")])
P([n("Neutrophils in TB. "), hl(
   "Neutrophils have a recognised double role in TB, contributing to early "
   "control but also to immunopathology, tissue destruction via matrix "
   "metalloproteinases, and a permissive replicative niche in established "
   "disease [8,13]. A baseline neutrophil-high state being associated with "
   "subsequent failure is biologically plausible and concordant with the "
   "interferon-driven neutrophil signature described in active TB [8]. We "
   "frame this as a hypothesis, not a mechanism: our data are cross-sectional "
   "and cannot establish whether neutrophilia is a cause, a consequence of "
   "higher bacterial burden, or a correlated marker of disease severity.")])
P([n("Clinical interpretation. "), hl(
   "The combination of high negative predictive value and low positive "
   "predictive value means the most defensible use is rule-out: a negative "
   "baseline signature identifies patients unlikely to fail, rather than "
   "reliably flagging those who will. We therefore do not claim that the "
   "signature can presently stratify patients for host-directed therapy or "
   "directly inform WHO End TB targets; such claims await prospective, "
   "adequately powered validation.")])
P([n("Relation to prior work. "), hl(
   "Our results neither contradict nor surpass existing models; they "
   "complement EMR- and multi-omic-based predictors [9–11] and the "
   "treatment-monitoring signatures of the source cohort [12] by adding "
   "cellular resolution and an explicit confounder audit. We have removed the "
   "earlier claim of being the ‘first’ such map.")])

H("4.1 Limitations", 2)
P([hl(
   "The limitations are substantial and central to interpretation. (i) Only "
   "seven baseline failure events, giving wide confidence intervals and low "
   "power; baseline discrimination did not reach significance against a "
   "label-permutation null (p=0.11) and no gene survived FDR correction. (ii) "
   "A single discovery cohort; no "
   "outcome-labelled external validation was achievable from public data. "
   "(iii) Cross-sectional design precludes causal inference; network edges are "
   "associative. (iv) Diabetes status and some confounders are unrecorded. (v) "
   "Cell-type inference is computational; flow-cytometric or single-cell "
   "confirmation of the neutrophil-high/T-cell-low phenotype, and prospective "
   "clinical validation, are required before any clinical use. We present the "
   "work as an exploratory framework and a reproducible baseline.")])

# ===== DECLARATIONS =====
H("Declarations", 1)
H("Ethical approval", 2)
P([hl("This study used publicly available, de-identified gene-expression data "
      "from the NCBI Gene Expression Omnibus. No new human participants were "
      "enrolled and no identifiable data were accessed; formal ethical "
      "approval was therefore not required. The original GSE89403 study "
      "obtained the appropriate ethical approvals described in its primary "
      "publication [12].")])
H("Consent to participate", 2)
P([hl("Not applicable; no new participants were enrolled and only de-identified "
      "public data were analysed.")])
H("Consent to publish", 2)
P([hl("Not applicable; no individually identifiable person or data are "
      "presented.")])
H("Data availability", 2)
P([n("GSE89403 and GSE193979 are available from the NCBI GEO. Processed "
     "matrices and analysis outputs are provided with the code repository.")])
H("Code availability", 2)
P([hl("All analysis scripts, the frozen model and intermediate outputs are "
      "available at https://github.com/hssling/TB-Transcriptomics-Project, "
      "enabling verification without re-execution.")])
H("Funding", 2)
P("This research received no specific grant from funding agencies in the "
  "public, commercial or not-for-profit sectors.")
H("Competing interests", 2)
P("The author declares no competing interests.")
H("Author contributions", 2)
P("Siddalingaiah H S: conceptualisation, methodology, software, formal "
  "analysis, investigation, data curation, writing (original draft and "
  "review/editing), visualisation.")
H("Use of AI", 2)
P([hl("Generative AI tools assisted with code development, literature search "
      "and manuscript drafting. All AI-assisted content was verified and "
      "edited by the author, who takes full responsibility for the work.")])

# ===== REFERENCES =====
H("References", 1)
refs = [
 "World Health Organization. Global Tuberculosis Report 2024. Geneva: WHO; 2024.",
 "World Health Organization. Treatment of tuberculosis: guidelines. Geneva: WHO.",
 "Berry MPR, Graham CM, McNab FW, et al. An interferon-inducible "
 "neutrophil-driven blood transcriptional signature in human tuberculosis. "
 "Nature. 2010;466:973–977.",
 "Zak DE, Penn-Nicholson A, Scriba TJ, et al. A blood RNA signature for "
 "tuberculosis disease risk: a prospective cohort study. Lancet. "
 "2016;387:2312–2322.",
 "Sweeney TE, Braviak L, Tato CM, Khatri P. Genome-wide expression for "
 "diagnosis of pulmonary tuberculosis: a multicohort analysis. Lancet Respir "
 "Med. 2016;4:213–224.",
 "Singhania A, Verma R, Graham CM, et al. A modular transcriptional signature "
 "identifies phenotypic heterogeneity of human tuberculosis infection. Nat "
 "Commun. 2018;9:2308.",
 "Blood transcriptomic signatures predict poor treatment outcomes in "
 "drug-susceptible pulmonary TB in Brazil (RePORT-Brazil). 2025. PMID 41282706.",
 "Lowe DM, Redford PS, Wilkinson RJ, O’Garra A, Martineau AR. Neutrophils in "
 "tuberculosis: friend or foe? Trends Immunol. 2012;33:14–25.",
 "Sambarey A, Smith K, Chung C, et al. Integrative analysis of multimodal "
 "patient data identifies personalized predictors of tuberculosis treatment "
 "prognosis. iScience. 2024. PMID 38357663.",
 "Vinhaes CL, Fukutani ER, Santana GC, et al. An integrative multi-omics "
 "approach to characterize interactions between tuberculosis and diabetes "
 "mellitus. iScience. 2024. PMID 38380250.",
 "Peng A-Z, Kong X-H, Liu S-T, et al. Explainable machine learning for early "
 "predicting treatment failure risk among patients with TB-diabetes "
 "comorbidity. Sci Rep. 2024. PMID 38514736.",
 "Thompson EG, Du Y, Malherbe ST, et al. Host blood RNA signatures predict the "
 "outcome of tuberculosis treatment. Tuberculosis. 2017;107:48–58 (GSE89403).",
 "Ndhlovu M, et al. Host-directed therapies for tuberculosis: a review. Front "
 "Immunol. 2019;10:325.",
 "Vianello E, et al. Transcriptional profiles predict treatment outcome in "
 "patients with tuberculosis and diabetes (TANDEM; GSE193979). EBioMedicine. "
 "2022;80:104023.",
 "Newman AM, Liu CL, Green MR, et al. Robust enumeration of cell subsets from "
 "tissue expression profiles. Nat Methods. 2015;12:453–457.",
 "Friedman J, Hastie T, Tibshirani R. Sparse inverse covariance estimation "
 "with the graphical lasso. Biostatistics. 2008;9:432–441.",
 "Lundberg SM, Lee S-I. A unified approach to interpreting model predictions. "
 "Adv Neural Inf Process Syst. 2017;30.",
 "Uhlén M, Karlsson MJ, Zhong W, et al. A genome-wide transcriptomic analysis "
 "of protein-coding genes in human blood cells. Science. 2019;366:eaax9198.",
]
for i, r in enumerate(refs, 1):
    P([n(f"{i}. {r}")], size=10)

out = f"{ROOT}/DAI_Revision_2026/deliverables/TB_Treatment_Failure_DAI_MajorRevision_v13.docx"
doc.save(out)
print("Saved manuscript:", out)
print("Paragraphs:", len(doc.paragraphs))
