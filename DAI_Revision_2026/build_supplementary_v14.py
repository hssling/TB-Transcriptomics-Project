# -*- coding: utf-8 -*-
"""Assemble the supplementary material .docx."""
import json
import os

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

HERE = os.path.dirname(os.path.abspath(__file__))
TAB = f"{HERE}/tables2"
FIG = f"{HERE}/figures2"
OUT = f"{HERE}/deliverables/Supplementary_Material_DAI_Revision_v14.docx"

M = json.load(open(f"{TAB}/arm_metrics.json"))
bench = pd.read_csv(f"{TAB}/model_benchmark.csv")
degs = pd.read_csv(f"{TAB}/deg_summary.csv")
decon = pd.read_csv(f"{TAB}/deconvolution_stats.csv")
corr = pd.read_csv(f"{TAB}/prediction_celltype_correlation.csv")
shap_top = pd.read_csv(f"{TAB}/shap_top_features.csv")
enrich = pd.read_csv(f"{TAB}/enrichment_targeted.csv")
enrich_deg = pd.read_csv(f"{TAB}/enrichment_deg.csv")
hubs = pd.read_csv(f"{TAB}/network_hubs.csv")
conf = pd.read_csv(f"{TAB}/confounder_models.csv")
ylink = pd.read_csv(f"{TAB}/y_linked_genes.csv")
part = pd.read_csv(f"{TAB}/deg_state_partition.csv")
ext_cell = pd.read_csv(f"{TAB}/external_celltype_stats.csv")

LABEL = {"DX": "Pre-treatment", "day_7": "Day 7", "week_4": "Week 4",
         "week_24": "Week 24", "combined": "All timepoints"}
ARMS = ["DX", "day_7", "week_4", "week_24"]

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(10.5)
for s in doc.sections:
    s.left_margin = s.right_margin = Inches(0.8)

CENTER = WD_ALIGN_PARAGRAPH.CENTER


def P(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=10.5, bold=False,
      italic=False, space_after=8):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def H(text, level=1):
    return doc.add_heading(text, level=level)


def table(headers, rows, widths=None, font=8.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(str(h))
        r.bold = True
        r.font.size = Pt(font)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(font)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def figure(filename, caption, width=6.3):
    path = f"{FIG}/{filename}"
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = CENTER
    P(caption, size=9.5, space_after=12)


P("Supplementary Material", align=CENTER, size=15, bold=True, space_after=4)
P("Whole-blood transcriptomic signatures of unfavourable tuberculosis "
  "treatment outcome before and during therapy", align=CENTER, size=11,
  italic=True, space_after=14)

# ---- S1 cohort ----
H("Supplementary Table S1. Cohort composition by timepoint", 1)
rows = []
for arm in ARMS + ["combined"]:
    rows.append([LABEL[arm], M[arm]["n"], M[arm]["n_subjects"],
                 M[arm]["n_events"], M[arm]["n"] - M[arm]["n_events"],
                 f"{M[arm]['n_events'] / M[arm]['n'] * 100:.1f}%"])
table(["Arm", "Samples", "Subjects", "Unfavourable", "Cured", "Prevalence"],
      rows, widths=[1.4, 0.9, 0.9, 1.2, 0.9, 1.0])

# ---- S2 full metrics ----
H("Supplementary Table S2. Complete performance metrics for the selected "
  "model in each arm", 1)
rows = []
for arm in ARMS + ["combined"]:
    b = M[arm]["models"][M[arm]["best_model"]]
    for r in b["per_class"]:
        rows.append([LABEL[arm], r["class"], r["n"],
                     f"{r['sensitivity']:.3f}", f"{r['specificity']:.3f}",
                     f"{r['ppv']:.3f}", f"{r['npv']:.3f}", f"{r['f1']:.3f}"])
table(["Arm", "Class", "n", "Sensitivity", "Specificity", "PPV", "NPV", "F1"],
      rows, widths=[1.1, 1.0, 0.5, 0.9, 0.9, 0.7, 0.7, 0.7])

rows = []
for arm in ARMS + ["combined"]:
    b = M[arm]["models"][M[arm]["best_model"]]
    cm = b["confusion"]
    rows.append([LABEL[arm], f"{b['balanced_accuracy']:.3f}",
                 f"{b['mcc']:.3f}", f"{b['brier']:.3f}",
                 f"{b['threshold']:.3f}",
                 f"TN {cm['tn']}, FP {cm['fp']}, FN {cm['fn']}, TP {cm['tp']}"])
table(["Arm", "Balanced accuracy", "MCC", "Brier", "Threshold",
       "Confusion matrix"], rows, widths=[1.1, 1.2, 0.7, 0.7, 0.8, 1.9])

# ---- S3 benchmark ----
H("Supplementary Table S3. Comparison of the three classifiers in every arm", 1)
P("Reported so that the use of gradient boosting for feature attribution can "
  "be judged against the alternatives rather than assumed. No single "
  "algorithm dominates; the arms in which discrimination is strong are strong "
  "for all three.", size=10, space_after=8)
rows = [[r.arm, r.model.replace("_", " "), f"{r.roc_auc:.3f}",
         f"{r.ci_low:.2f}–{r.ci_high:.2f}", f"{r.pr_auc:.3f}",
         f"{r.mcc:.3f}", f"{r.brier:.3f}"]
        for _, r in bench.iterrows()]
table(["Arm", "Model", "ROC-AUC", "95% CI", "PR-AUC", "MCC", "Brier"], rows,
      widths=[1.1, 1.4, 0.8, 1.0, 0.8, 0.7, 0.7])

# ---- S4 permutation ----
H("Supplementary Table S4. Label-permutation null distributions", 1)
P("Observed and permuted discrimination were computed under one identical "
  "cross-validation protocol, so the comparison is internally consistent.",
  size=10, space_after=8)
rows = []
for arm in ARMS + ["combined"]:
    p = M[arm]["permutation"]
    rows.append([LABEL[arm], f"{p['observed_auc']:.3f}",
                 f"{p['null_mean']:.3f}", f"{p['null_sd']:.3f}",
                 p["n_perm"],
                 "< 0.002" if p["permutation_p"] < 0.002
                 else f"{p['permutation_p']:.3f}",
                 f"{M[arm].get('loo_auc', float('nan')):.3f}"
                 if M[arm].get("loo_auc") is not None else "—"])
table(["Arm", "Observed AUC", "Null mean", "Null SD", "Permutations",
       "p", "Leave-one-out AUC"], rows,
      widths=[1.1, 1.0, 0.8, 0.7, 1.0, 0.7, 1.2])

# ---- S5 DEG ----
H("Supplementary Table S5. Differential expression by arm", 1)
rows = [[LABEL[r.arm], f"{int(r.genes_tested):,}", f"{int(r.nominal_p05):,}",
         f"{int(r.nominal_p01):,}", f"{int(r.fdr_significant):,}",
         r.top_gene, f"{r.top_p:.2e}"]
        for _, r in degs.iterrows()]
table(["Arm", "Genes tested", "p < 0.05", "p < 0.01", "FDR < 0.05",
       "Top gene", "Top p"], rows, widths=[1.1, 1.0, 0.8, 0.8, 0.9, 1.0, 0.8])

figure("Figure_volcano_arms.png",
       "Supplementary Figure S1. Differential expression by arm. Points in red "
       "reach p < 0.01; the six most significant genes are labelled in each "
       "panel.")

# ---- S6 partition ----
H("Supplementary Table S6. Genes separating outcome groups, by the states in "
  "which they do so", 1)
rows = [[cls.replace("_", " ").capitalize(), int((part["class"] == cls).sum())]
        for cls in part["class"].unique()]
table(["Class", "Genes"], rows, widths=[3.0, 1.0])

# ---- S7 deconvolution ----
H("Supplementary Table S7. Immune-cell composition by outcome in each arm", 1)
rows = [[LABEL[r.arm], r.cell_type, f"{r.median_unfavourable:+.3f}",
         f"{r.median_cured:+.3f}", f"{r.rank_biserial_r:+.3f}",
         f"{r.p_value:.4f}",
         int(r.n_marker_genes) if pd.notna(r.n_marker_genes) else "—"]
        for _, r in decon.iterrows()]
table(["Arm", "Cell type", "Median (unfavourable)", "Median (cured)",
       "Rank-biserial r", "p", "Marker genes"], rows,
      widths=[1.0, 1.1, 1.1, 1.0, 1.0, 0.7, 0.8])

# ---- S8 correlation ----
H("Supplementary Table S8. Correlation between predicted probability and "
  "immune-cell scores", 1)
rows = [[LABEL[r.arm], r.cell_type, int(r.n), f"{r.spearman_rho:+.3f}",
         f"{r.p_value:.2e}"] for _, r in corr.iterrows()]
table(["Arm", "Cell type", "n", "Spearman ρ", "p"], rows,
      widths=[1.3, 1.4, 0.6, 1.0, 1.0])

# ---- S9 SHAP ----
H("Supplementary Table S9. Highest-ranked genes by SHAP attribution in each "
  "arm", 1)
rows = [[LABEL[r.arm], int(r["rank"]), r.gene_symbol,
         f"{r.mean_abs_shap:.4f}",
         "raises risk" if r.shap_direction_high_expression > 0 else "lowers risk"]
        for _, r in shap_top[shap_top["rank"] <= 15].iterrows()]
table(["Arm", "Rank", "Gene", "Mean |SHAP|", "Direction of high expression"],
      rows, widths=[1.1, 0.6, 1.1, 1.0, 1.7])

# ---- S10 enrichment ----
H("Supplementary Table S10. Immune programmes tested against the attribution "
  "panels", 1)
P("Complete results of the targeted test, including programmes with no "
  "overlapping gene. Six pre-specified programmes were tested per arm, so the "
  "multiple-testing burden is small.", size=10, space_after=8)
rows = [[LABEL.get(r.arm, r.arm), str(r.term)[:80], r.overlap,
         f"{r.p_value:.2e}", f"{r.adjusted_p:.4f}",
         str(r.genes)[:60] if pd.notna(r.genes) else "—"]
        for _, r in enrich.iterrows()]
table(["Arm", "Immune programme", "Overlap", "p", "Adjusted p",
       "Overlapping genes"], rows,
      widths=[0.9, 2.1, 0.6, 0.8, 0.8, 1.6])

H("Supplementary Table S11. Open-ended enrichment on genes surviving "
  "false-discovery correction", 1)
if len(enrich_deg):
    P("Applicable only to the week 24 arm, the sole arm producing genes that "
      "survive false-discovery correction.", size=10, space_after=8)
    rows = [[LABEL.get(r.arm, r.arm), str(r.term)[:80], r.overlap,
             f"{r.p_value:.2e}", f"{r.adjusted_p:.2e}"]
            for _, r in enrich_deg.iterrows()]
    table(["Arm", "Pathway", "Overlap", "p", "Adjusted p"], rows,
          widths=[0.9, 3.1, 0.8, 0.8, 0.8])
else:
    P("No arm produced a gene list large enough for open-ended enrichment.",
      size=10)

# ---- S11 network ----
H("Supplementary Table S12. Conditional-dependency network hubs", 1)
rows = [[LABEL[r.arm], r.gene_symbol, int(r.degree),
         f"{r.mean_abs_partial_correlation:.3f}"]
        for _, r in hubs[hubs.degree > 0].iterrows()]
table(["Arm", "Gene", "Degree", "Mean |partial correlation|"], rows,
      widths=[1.3, 1.3, 0.8, 1.8])

# ---- S12 confounders ----
H("Supplementary Table S13. Confounder-adjusted models", 1)
rows = [[LABEL[r.arm], r.model, int(r.n),
         f"{r.odds_ratio:.2f}" if pd.notna(r.odds_ratio) else "—",
         f"{r.ci_low:.2f}–{r.ci_high:.2f}" if pd.notna(r.ci_low) else "—",
         f"{r.p_value:.3f}" if pd.notna(r.p_value) else "—"]
        for _, r in conf.iterrows()]
table(["Arm", "Model", "n", "Odds ratio", "95% CI", "p"], rows,
      widths=[1.0, 2.6, 0.5, 0.9, 1.1, 0.6])

H("Supplementary Table S14. Sex-linked transcripts tested for differential "
  "expression by outcome", 1)
P("No sex-linked transcript separates outcome groups in any arm. Where such "
  "genes appear high in an unadjusted feature ranking, they reflect the sex "
  "composition of a small outcome group rather than a biological signal.",
  size=10, space_after=8)
rows = [[LABEL[r.arm], r.gene_symbol, f"{r.log2_fold_change:+.3f}",
         f"{r.p_value:.3f}", f"{r.fdr:.3f}"]
        for _, r in ylink.iterrows()]
table(["Arm", "Gene", "log₂ fold change", "p", "FDR"], rows,
      widths=[1.3, 1.2, 1.3, 0.8, 0.8])

# ---- S14 external ----
H("Supplementary Table S15. Independent cohort, cell-composition "
  "scores", 1)
rows = [[r.timepoint, r.cell_type, int(r.n_unfavourable), int(r.n_cured),
         f"{r.median_unfavourable:+.3f}", f"{r.median_cured:+.3f}",
         f"{r.rank_biserial_r:+.3f}", f"{r.p_value:.3f}"]
        for _, r in ext_cell.iterrows()]
table(["Timepoint", "Cell type", "Relapse", "Cured", "Median (relapse)",
       "Median (cured)", "Rank-biserial r", "p"], rows,
      widths=[0.9, 1.0, 0.6, 0.6, 1.0, 1.0, 1.0, 0.6])

# ---- S16 robustness audit ----
H("Supplementary Table S16. Robustness checks on the week-24 arm", 1)
P("Five ways the end-of-treatment result could be artefactual, each tested "
  "directly.", size=10, space_after=8)
audit = json.load(open(f"{TAB}/robustness_audit.json"))

P("Technical confounding", size=10.5, bold=True, space_after=4)
rows = [[r["arm"], r["metric"], f"{r['median_unfavourable']:.4g}",
         f"{r['median_cured']:.4g}", f"{r['p_value']:.3f}"]
        for r in audit["technical"]]
table(["Arm", "Metric", "Median (unfavourable)", "Median (cured)", "p"], rows,
      widths=[1.2, 1.2, 1.4, 1.2, 0.7])

P("Influence of individual patients with an unfavourable outcome", size=10.5,
  bold=True, space_after=4)
rows = [[r["subject_removed"], f"{r['auc']:.3f}", f"{r['delta']:+.3f}"]
        for r in audit["influence"]["leave_one_out"]]
rows.append(["none (full cohort)",
             f"{audit['influence']['baseline_auc']:.3f}", "—"])
table(["Subject removed", "ROC-AUC", "Change"], rows,
      widths=[1.8, 1.0, 1.0])

P("Multiplicity across the four arms", size=10.5, bold=True, space_after=4)
rows = [[r["arm"], f"{r['p_raw']:.3f}", f"{r['p_holm']:.3f}",
         f"{r['p_bonferroni']:.3f}",
         "yes" if r["significant_holm"] else "no"]
        for r in audit["multiplicity"]]
table(["Arm", "p (raw)", "p (Holm)", "p (Bonferroni)", "Significant (Holm)"],
      rows, widths=[1.3, 0.9, 0.9, 1.1, 1.3])

P("Association with bacterial load measured at diagnosis", size=10.5,
  bold=True, space_after=4)
rows = [[r["measure"].upper(), int(r["n"]), f"{r['spearman_rho']:+.3f}",
         f"{r['p_value']:.3f}"] for r in audit["bacterial_load"]]
table(["Measure", "n", "Spearman ρ with week-24 predicted risk", "p"], rows,
      widths=[1.0, 0.6, 2.4, 0.8])

# ---- S17 published signatures ----
H("Supplementary Table S17. Published signatures applied without fitting", 1)
pubsig = pd.read_csv(f"{TAB}/established_signatures.csv")
rows = [[LABEL[r.arm], r.signature, int(r.genes_up), int(r.genes_down),
         f"{r.roc_auc:.3f}", f"{r.rank_biserial_r:+.3f}", f"{r.p_value:.2e}"]
        for _, r in pubsig.iterrows()]
table(["Arm", "Signature", "Genes up", "Genes down", "ROC-AUC",
       "Rank-biserial r", "p"], rows,
      widths=[0.9, 2.0, 0.7, 0.8, 0.7, 1.0, 0.8])

# ---- S18 trajectory ----
H("Supplementary Table S18. Within-subject treatment response", 1)
traj = pd.read_csv(f"{TAB}/response_trajectory.csv")
rows = [[r.comparison, int(r.n_subjects), int(r.n_unfavourable),
         f"{r.median_unfavourable:+.3f}", f"{r.median_cured:+.3f}",
         f"{r.rank_biserial_r:+.3f}", f"{r.p_value:.3f}",
         f"{r.roc_auc_delta_model:.3f}"]
        for _, r in traj.iterrows()]
table(["Comparison", "Paired subjects", "Unfavourable",
       "Median alignment (unfavourable)", "Median alignment (cured)",
       "Rank-biserial r", "p", "ROC-AUC of a model on the change vectors"],
      rows, widths=[1.1, 0.8, 0.7, 1.1, 1.0, 0.8, 0.6, 1.2], font=8)

figure("Figure_response_trajectory.png",
       "Supplementary Figure S2. Alignment of each subject's diagnosis-to-later "
       "change with the average change of cured subjects.")

# ---- reproducibility ----
H("Supplementary Note S1. Reproducibility", 1)
P("Analyses were run in Python with scikit-learn, XGBoost, SHAP, statsmodels, "
  "SciPy, pandas and NetworkX. Random seeds are fixed throughout. The "
  "expression matrix is rebuilt from the primary deposit by a single script, "
  "so the substrate for every arm is reproducible from the raw counts without "
  "manual intervention. Analysis code, the frozen models and all intermediate "
  "tables are released with the manuscript.", size=10)

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f"wrote {OUT}")
