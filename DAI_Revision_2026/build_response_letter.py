# -*- coding: utf-8 -*-
"""Build the point-by-point Response-to-Reviewers letter as a table (.docx).
Every editor / R1 / R2 / R3 comment -> response -> manuscript location."""
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = "d:/research-automation/TB multiomics/TB-Treatment-Failure-Clean"
A = json.load(open(f"{ROOT}/DAI_Revision_2026/tables/wpA_metric_suite.json"))
rf = A["models"]["RandomForest"]
doc = Document()
doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(10)

h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = h.add_run("Point-by-Point Response to Reviewers"); r.bold = True; r.font.size = Pt(15)
doc.add_paragraph("Manuscript: “Baseline whole-blood transcriptomic risk "
                  "stratification for unfavourable tuberculosis treatment outcome” "
                  "(Submission 28df71d5-9f1a-4e4d-8cd3-8fcf4ef17dfb). Discover "
                  "Artificial Intelligence.")
p = doc.add_paragraph()
p.add_run("We thank the Editor and three reviewers for rigorous, constructive "
          "criticism. We have comprehensively revised the manuscript: removing all "
          "causal/‘first map’ language, reframing the work as an exploratory, "
          "associative BASELINE risk-stratification study, rebuilding every analysis "
          "on a leakage-free pre-treatment cohort, and reporting results with full "
          "transparency including their limitations. Below, each comment is answered "
          "individually with the manuscript location. Revised text is highlighted in "
          "the manuscript.").italic = True


def block(title):
    p = doc.add_paragraph(); rr = p.add_run(title); rr.bold = True; rr.font.size = Pt(12)
    rr.font.color.rgb = RGBColor(0x1F, 0x3B, 0x73)


def qa_table(items):
    t = doc.add_table(rows=1, cols=3); t.style = "Light Grid Accent 1"
    hdr = ["Reviewer comment", "Author response", "Location"]
    for j, c in enumerate(hdr):
        cell = t.rows[0].cells[j]; cell.paragraphs[0].add_run(c).bold = True
    widths = [0.34, 0.52, 0.14]
    for q, a, loc in items:
        cells = t.add_row().cells
        cells[0].text = q; cells[1].text = a; cells[2].text = loc
    doc.add_paragraph()


AUC = f"{rf['roc_auc']:.2f} (95% CI {rf['roc_auc_ci95'][0]:.2f}–{rf['roc_auc_ci95'][1]:.2f})"

block("Editor")
qa_table([
 ("Response-to-reviewers file (point-by-point).",
  "Provided (this document), with per-comment responses and locations.", "This file"),
 ("Manuscript with tracked changes / highlighting (mandatory).",
  "All revised/added content is highlighted in yellow in the revised manuscript.",
  "Throughout"),
 ("Add a 'Declarations' section with Ethical approval, Consent to participate, "
  "Consent to publish as separate subheadings.",
  "Added a dedicated Declarations section with these three subheadings addressed "
  "individually, plus data/code availability, funding, competing interests.",
  "Declarations"),
 ("Concerns re novelty, methodology, experimental evaluation, statistical analysis.",
  "Addressed across the revision: novelty reframed and prior work cited (R3.1); "
  "leakage-free baseline methodology; full metric suite with CIs; effect sizes and "
  "FDR throughout.", "Sec 2–3"),
])

block("Reviewer 1")
qa_table([
 ("Independent validation of ML performance in ≥2 separate studies; "
  "cross-validation is weak and signatures don't generalize.",
  "We agree generalisation is essential. We rebuilt evaluation on leakage-free "
  "pre-treatment samples with repeated stratified CV, leave-one-out CV and "
  "bootstrap CIs, and honestly report modest discrimination (ROC-AUC " + AUC + "). "
  "A genuine GEO search identified no usable outcome-labelled independent cohort "
  "(closest, GSE193979, lacks public per-patient outcomes and a public sample "
  "bridge); we document this in a Supplementary exclusion table (per Reviewer 2's "
  "option) and provide a label-free external portability check (ρ=−0.41). We "
  "also temper all generalisation claims.", "Sec 2.3, 3.2, 3.8; Table S5"),
 ("Writing reads like a student report; lacks context, data description, rationale; "
  "abrupt transitions; how were mechanisms/features chosen.",
  "The manuscript was substantially rewritten: expanded Introduction with prior "
  "signatures and gaps; a full Methods describing the dataset, inclusion/exclusion, "
  "sample/feature counts and preprocessing; and Results that motivate each step. "
  "Feature selection (univariate-in-fold + SHAP) and the path from prediction to "
  "cellular interpretation are now explicit.", "Sec 1–3"),
 ("Benchmark/compare your ML model against similar approaches in literature.",
  "Added a benchmark table (Table S6) comparing against Thompson 2017, TANDEM 2022, "
  "RePORT-Brazil 2025 and the multimodal/EMR/multi-omic ML models (PMIDs 38357663, "
  "38380250, 38514736). We position our contribution as cellular resolution and "
  "confounder transparency, not predictive superiority.", "Sec 3.8; Table S6"),
 ("Define DOTS in the Introduction.",
  "DOTS is now defined at first use ('Directly Observed Treatment, Short-course').",
  "Sec 1"),
 ("Specify what has been accomplished in literature on the Resolution and Logic gaps.",
  "The Introduction now cites prior bulk-signature work and its resolution "
  "limitation explicitly; we dropped the 'Logic gap'/causal framing entirely and "
  "replaced it with an associative, confounder-aware framing.", "Sec 1"),
 ("What features were used for training the ML models? Provide details.",
  "Methods now specify 16,147 gene features with univariate selection inside each "
  "fold; the full ranked predictor list with HGNC symbols is in Table S2.",
  "Sec 2.3; Table S2"),
 ("ROC-AUC in text conflicts with Figure 1.",
  "Resolved: all performance numbers are regenerated from a single pipeline and "
  "Figure 1 and the text now report the same values (ROC-AUC " + AUC + ").",
  "Fig 1; Sec 3.2"),
 ("Figure 2: what do 0 and 1 correspond to — success or failure?",
  "Figure legends now state explicitly: 0 = cure, 1 = treatment failure.",
  "Fig 1–2 legends"),
 ("Policy: how can patients be stratified for transcriptomic failure risk at the "
  "start, if these signatures appear later in treatment?",
  "An excellent point that we now make central. We restrict to BASELINE "
  "(pre-treatment) samples and show baseline discrimination is only modest; we "
  "demonstrate quantitatively that the stronger signal in the original submission "
  "was a timepoint/leakage artefact. Policy claims are removed.", "Sec 3.2; Disc"),
])

block("Reviewer 2")
qa_table([
 ("Class distribution must appear prominently in main Methods.",
  "Class distribution (7 failure vs 83 cure; prevalence 7.8%) is now stated in the "
  "main Methods and Results, with explicit note that it governs all metrics.",
  "Sec 2.1, 3.1"),
 ("AUC alone insufficient: report precision, recall, F1, confusion matrix and a "
  "Precision-Recall curve at the chosen operating threshold.",
  "All reported: PR curve (Fig 1B), confusion matrix (Fig 1D), and "
  "sensitivity/specificity/PPV/NPV/F1/MCC at the Youden operating point.",
  "Fig 1; Sec 3.2; Table S1"),
 ("External validation absent; validate on an independent GEO cohort, or document "
  "no suitable dataset with an exclusion table; deposit model on GitHub.",
  "We took the exclusion-table option you offered: a documented GEO search "
  "(Table S5) shows no usable outcome-labelled cohort; we add a label-free external "
  "portability check and deposit the frozen model and code on GitHub.",
  "Sec 3.8; Table S5; Code availability"),
 ("Full RFE feature list, both signatures, and full Supplementary Table 1 are "
  "inaccessible — findings unverifiable.",
  "The full ranked feature list with symbols (Table S2), both cell-type signatures "
  "(Methods 2.5; Table S8), and complete result tables are now provided, plus a "
  "public repository for verification.", "Tables S2, S8"),
 ("SHAP-based interpretation and formal DEG with volcano and pathway enrichment "
  "strongly recommended.",
  "Added: SHAP summary/bar (Methods 2.4), formal Mann-Whitney DEG with true-log2 "
  "volcano (Fig S1; Table S3) and enrichment (Table S10).", "Sec 2.4, 3.3"),
 ("PBMC3k lacks neutrophils and is inappropriate for neutrophil-signature "
  "validation; use whole-blood scRNA.",
  "We removed the PBMC3k neutrophil validation and instead validate cell-type "
  "specificity in the granulocyte-containing Human Protein Atlas blood atlas; "
  "10/12 neutrophil-signature genes are neutrophil-specific.", "Sec 2.6, 3.5; Fig 3"),
 ("Alignment between CIBERSORT deconvolution and ML findings asserted but never "
  "quantified.",
  "Now quantified: the model's failure probability correlates with the deconvolved "
  "neutrophil score (Spearman ρ=0.66) and inversely with the T-cell score "
  "(ρ=−0.66), p≈1e-12.", "Sec 3.4"),
 ("Remove all causal language for Glasso hubs; GGM edges are undirected conditional "
  "associations and the outcome is not an input to Glasso.",
  "Done. We describe the network strictly as undirected conditional associations; "
  "hubs are 'hub genes', not regulators/targets; we state the outcome is not an "
  "input.", "Sec 2.7, 3.7"),
 ("Neutrophil-high/T-cell-low may be reactive/confounded; HIV, diabetes, bacterial "
  "load, drug resistance unaddressed.",
  "Added a confounder audit: HIV and drug resistance were exclusion criteria in the "
  "source cohort; bacterial load (MGIT/Xpert) is adjusted for and partially "
  "attenuates the association; sex is adjusted and is non-significant. Diabetes is "
  "unrecorded (stated as a limitation).", "Sec 2.8, 3.6; Table S7"),
 ("Foreground baseline patient risk stratification as the defensible application.",
  "This is now the central framing of the paper, from title to conclusions.",
  "Throughout"),
 ("Temper claims of directly informing the WHO End TB Strategy.",
  "All WHO End TB / clinical-readiness claims removed; we explicitly state the "
  "signature is not yet clinically actionable.", "Disc 4"),
 ("Mann-Whitney results need violin plots, distributional assessment, and "
  "rank-biserial effect sizes alongside p-values.",
  "Added violin plots with individual points and significance annotations (Fig 2) "
  "and rank-biserial effect sizes alongside every p-value (Table S4).",
  "Fig 2; Table S4"),
])

block("Reviewer 3")
qa_table([
 ("'First multi-omic mechanistic map' is incorrect (PMIDs 38357663, 38380250, "
  "38514736); compare V2 against existing methods.",
  "We retract the 'first' claim, cite all three works, and add a benchmark table "
  "positioning our contribution (baseline cellular-deconvolution + confounder "
  "audit) relative to them.", "Sec 1, 3.8; Table S6"),
 ("Insufficient data detail: datasets, processing, inclusion/exclusion, sample and "
  "feature counts, meaning of 'standardized pipeline'.",
  "Methods now give the cohort, inclusion/exclusion, N=90 samples and 16,147 "
  "features, and define 'standardised pipeline' as the scripted normalisation/"
  "transform/harmonisation sequence released as code.", "Sec 2.1–2.2"),
 ("Feature-importance method not elaborated.",
  "Methods now describe univariate-in-fold selection and SHAP TreeExplainer; full "
  "list in Table S2.", "Sec 2.4; Table S2"),
 ("Evidence not shown for 'Y-linked genes top predictors' and '50 genes correlated'.",
  "Now shown: the DEG table (Table S3) and a dedicated analysis demonstrate ~1,600 "
  "nominally associated genes (0 surviving FDR) and that Y-linked genes are NOT "
  "differentially expressed at baseline — they were a sex-confound artefact, now "
  "explained.", "Sec 3.3, 3.6; Table S3"),
 ("Neutrophil-high/T-cell-low claim rests on model interpretation; in vitro/in vivo "
  "validation should be conducted.",
  "As a computational study we cannot perform wet-lab work; we instead (i) validate "
  "cell-type specificity in a granulocyte-containing reference, (ii) quantify "
  "deconvolution–model concordance, and (iii) list flow/single-cell confirmation as "
  "an explicit, necessary future step. We also temper the claim (T-cell reduction "
  "is non-significant at baseline).", "Sec 3.4–3.5, 4.1"),
 ("Single-cell dataset (Nathan 2021) cited to validate — can the author do this?",
  "We performed our own cell-type-specificity analysis in the HPA blood atlas "
  "(includes neutrophils, unlike PBMC3k), reported quantitatively in Fig 3.",
  "Sec 3.5; Fig 3"),
 ("References missing for several statements (MDR-TB cost, male neutrophilia, the "
  "'Trojan Horse' section, etc.).",
  "Unreferenced rhetorical passages were removed or rewritten with citations; the "
  "speculative 'Trojan Horse' section was deleted and replaced by a referenced, "
  "hedged neutrophil paragraph.", "Sec 1, 4"),
 ("Minor: missing in-text figure references.",
  "Every figure is now cited in the text at first mention.", "Throughout"),
 ("Minor: terms introduced without context (DOTS, SMOTE, signatures).",
  "Defined at first use; SMOTE was removed (we use class weighting, now stated).",
  "Sec 1–2"),
 ("Minor: Figure 2 include significance annotations.",
  "Added significance brackets/stars and effect sizes to the violin figure.",
  "Fig 2"),
 ("Minor: Figure 3 lacks cell-type labels and a defined scale.",
  "The new Fig 3 has labelled blood-cell types and a defined colour scale "
  "(row-max-normalised nTPM).", "Fig 3"),
 ("Minor: all figures hard to read; increase font size.",
  "All figures regenerated at 300 dpi with enlarged fonts and axis labels.",
  "All figures"),
])

out = f"{ROOT}/DAI_Revision_2026/deliverables/Response_to_Reviewers_DAI_MajorRevision.docx"
doc.save(out)
print("Saved response letter:", out)
